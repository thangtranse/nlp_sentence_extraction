from __future__ import annotations

import hashlib
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
WORK = ROOT / "artifacts" / "report_work"
REFERENCE = Path("/Users/thangtran/Downloads/cau-truc-tieu-luan-hutech.docx")
OUTPUT = ROOT / "artifacts" / "bao-cao-tom-tat-van-ban-textrank-hutech.docx"
EXPECTED_SHA = "102b60260797a88ca56026ca5c81380a663b32812da33bba02e5cdbb8941a5dd"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=110, bottom=80, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="A6A6A6", size="5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, separate, text, end):
        run._r.append(el)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def configure_styles(doc: Document) -> None:
    style_names = [s.name for s in doc.styles]
    normal = doc.styles["Normal"] if "Normal" in style_names else doc.styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.0)

    for name, size, before, after, color in (
        ("Heading 1", 15, 18, 12, "000000"),
        ("Heading 2", 14, 14, 8, "000000"),
        ("Heading 3", 13, 10, 6, "000000"),
    ):
        style_id = name.replace(" ", "")
        style = doc.styles.get_by_id(style_id, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)
    if "List Bullet" not in [s.name for s in doc.styles]:
        list_bullet = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        list_bullet.base_style = normal
    if "Table Grid" not in [s.name for s in doc.styles]:
        doc.styles.add_style("Table Grid", WD_STYLE_TYPE.TABLE)

    if "Caption Academic" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("Caption Academic", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Caption Academic"]
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cap.font.size = Pt(11)
    cap.font.italic = True
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(8)
    cap.paragraph_format.first_line_indent = Cm(0)

    if "Code Academic" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("Code Academic", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Academic"]
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.6)
    code.paragraph_format.right_indent = Cm(0.3)
    code.paragraph_format.first_line_indent = Cm(0)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0


def add_heading(doc: Document, text: str, level: int):
    style = doc.styles.get_by_id(f"Heading{level}", WD_STYLE_TYPE.PARAGRAPH)
    return doc.add_paragraph(text, style=style)


def add_body(doc: Document, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.35
        p.style.font.name = "Times New Roman"
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(cell)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cells[idx])
            for p in cells[idx].paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10.2)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_caption(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Caption Academic")


def make_figures() -> dict[str, Path]:
    WORK.mkdir(parents=True, exist_ok=True)
    figures = {}
    font_path = "/System/Library/Fonts/Supplemental/Arial.ttf"
    bold_path = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
    font = ImageFont.truetype(font_path, 23)
    small = ImageFont.truetype(font_path, 20)
    bold = ImageFont.truetype(bold_path, 23)

    img = Image.new("RGB", (2200, 520), "white")
    draw = ImageDraw.Draw(img)
    labels = ["Tài liệu", "Tách câu &\ntiền xử lý", "TF–IDF", "Cosine\nsimilarity", "Đồ thị câu", "PageRank", "Chọn theo\nbudget", "Khôi phục\nthứ tự"]
    box_w, box_h, gap, x0, y0 = 220, 125, 48, 35, 180
    for i, label in enumerate(labels):
        x = x0 + i * (box_w + gap)
        fill = "#FCE4D6" if i in (4, 5) else "#D9EAF7"
        draw.rounded_rectangle((x, y0, x + box_w, y0 + box_h), radius=18, fill=fill, outline="#315B7D", width=4)
        draw.multiline_text((x + box_w / 2, y0 + box_h / 2), label, font=font, fill="#102A43", anchor="mm", align="center", spacing=4)
        if i < len(labels) - 1:
            a, b, cy = x + box_w + 7, x + box_w + gap - 7, y0 + box_h / 2
            draw.line((a, cy, b, cy), fill="#315B7D", width=5)
            draw.polygon([(b, cy), (b - 15, cy - 10), (b - 15, cy + 10)], fill="#315B7D")
    p = WORK / "pipeline.png"
    img.save(p, dpi=(220, 220))
    figures["pipeline"] = p

    matrix = [
        [0.00, 0.62, 0.18, 0.05, 0.00],
        [0.62, 0.00, 0.41, 0.09, 0.03],
        [0.18, 0.41, 0.00, 0.55, 0.14],
        [0.05, 0.09, 0.55, 0.00, 0.47],
        [0.00, 0.03, 0.14, 0.47, 0.00],
    ]
    img = Image.new("RGB", (1100, 960), "white")
    draw = ImageDraw.Draw(img)
    draw.text((550, 55), "Minh họa ma trận tương đồng cosine", font=bold, fill="#102A43", anchor="mm")
    cell, ox, oy = 135, 210, 170
    for i in range(5):
        draw.text((ox - 55, oy + i * cell + cell / 2), f"S{i+1}", font=bold, fill="#102A43", anchor="mm")
        draw.text((ox + i * cell + cell / 2, oy - 45), f"S{i+1}", font=bold, fill="#102A43", anchor="mm")
        for j in range(5):
            val = matrix[i][j]
            shade = int(245 - val / 0.7 * 130)
            fill = (shade, min(250, shade + 20), 250)
            x1, y1 = ox + j * cell, oy + i * cell
            draw.rectangle((x1, y1, x1 + cell, y1 + cell), fill=fill, outline="#6F8FA8", width=2)
            draw.text((x1 + cell / 2, y1 + cell / 2), f"{val:.2f}", font=small, fill="#102A43", anchor="mm")
    p = WORK / "similarity_matrix.png"
    img.save(p, dpi=(220, 220))
    figures["matrix"] = p
    img = Image.new("RGB", (1300, 920), "white")
    draw = ImageDraw.Draw(img)
    positions = [(210, 260), (590, 165), (850, 430), (560, 720), (1080, 690)]
    scores = [0.17, 0.26, 0.24, 0.21, 0.12]
    for i in range(5):
        for j in range(i + 1, 5):
            val = matrix[i][j]
            if val >= 0.10:
                draw.line((positions[i], positions[j]), fill="#7A9BB8", width=max(2, int(val * 14)))
                mx = (positions[i][0] + positions[j][0]) / 2
                my = (positions[i][1] + positions[j][1]) / 2
                draw.text((mx, my), f"{val:.2f}", font=small, fill="#315B7D", anchor="mm", stroke_width=3, stroke_fill="white")
    for i, ((x, y), score) in enumerate(zip(positions, scores)):
        radius = int(70 + score * 120)
        draw.ellipse((x-radius, y-radius, x+radius, y+radius), fill="#D9EAF7", outline="#315B7D", width=5)
        draw.multiline_text((x, y), f"S{i+1}\nPR={score:.2f}", font=bold, fill="#102A43", anchor="mm", align="center", spacing=4)
    p = WORK / "sentence_graph.png"
    img.save(p, dpi=(220, 220))
    figures["graph"] = p
    return figures


def replace_cover_topic(doc: Document) -> None:
    target = "Nghiên cứu về việc áp dụng các mô hình học sâu/transformer trong việc"
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == target:
            for p2 in doc.paragraphs[i:i + 3]:
                p2.text = ""
            p.text = "TÓM TẮT VĂN BẢN TRÍCH XUẤT BẰNG TEXTRANK/PAGERANK VÀ MMR"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.name = "Times New Roman"
            p.runs[0].font.size = Pt(15)
            doc.paragraphs[i + 1].text = "TRÊN ĐỒ THỊ CÂU CÓ TRỌNG SỐ"
            doc.paragraphs[i + 1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[i + 1].runs[0].bold = True
            doc.paragraphs[i + 1].runs[0].font.name = "Times New Roman"
            doc.paragraphs[i + 1].runs[0].font.size = Pt(15)
            break
    for p in doc.paragraphs:
        if p.text.strip().startswith("Thành phố Hồ Chí Minh, tháng"):
            p.text = "Thành phố Hồ Chí Minh, tháng 8 năm 2026"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def rebuild_front_lists(doc: Document) -> None:
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip() == "MỤC LỤC":
            # Use the following blank paragraph as a live Word TOC field.
            if i + 1 < len(paras):
                paras[i + 1].text = ""
                add_field(paras[i + 1], 'TOC \\o "1-3" \\h \\z \\u', "Nhấn Ctrl+A, F9 để cập nhật mục lục")
                paras[i + 1].paragraph_format.first_line_indent = Cm(0)
                paras[i + 1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.text.strip() == "DANH MỤC TỪ VIẾT TẮT" and i + 1 < len(paras):
            paras[i + 1].text = "DUC: Document Understanding Conference\nNLP: Natural Language Processing\nTF–IDF: Term Frequency–Inverse Document Frequency\nPR: PageRank\nMMR: Maximal Marginal Relevance\nP@K: Precision at K\nR@K: Recall at K\nF1@K: F1-score at K"
            paras[i + 1].paragraph_format.first_line_indent = Cm(0)
            paras[i + 1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.text.strip() == "DANH MỤC HÌNH ẢNH" and i + 1 < len(paras):
            paras[i + 1].text = "Hình 3.1. Luồng xử lý tóm tắt trích xuất\nHình 4.1. Chất lượng khi tinh chỉnh ngưỡng ở vùng thấp\nHình 4.2. Sức khỏe đồ thị khi tinh chỉnh ngưỡng"
            paras[i + 1].paragraph_format.first_line_indent = Cm(0)
            paras[i + 1].alignment = WD_ALIGN_PARAGRAPH.LEFT
        if p.text.strip() == "DANH MỤC BẢNG" and i + 1 < len(paras):
            paras[i + 1].text = "Bảng 1.1. Cấu trúc dữ liệu đầu vào và đầu ra\nBảng 2.1. Dữ liệu kiểm tra TF–IDF viết tay\nBảng 2.2. Kết quả TF–IDF trước và sau chuẩn hóa L2\nBảng 2.3. Đối chiếu TF–IDF viết tay với scikit-learn\nBảng 2.4. Nhóm chỉ số đánh giá trong notebook\nBảng 3.1. Cấu trúc package src/nlp_practice\nBảng 3.2. Tham số của cấu hình cuối\nBảng 4.1. Thiết kế thực nghiệm\nBảng 4.2. Sweep thô ngưỡng similarity\nBảng 4.3. Sweep mịn PageRank + Top-K=15\nBảng 4.4. Xác nhận MMR/100 từ\nBảng 4.5. So sánh các cấu hình MMR\nBảng 4.6. Kết quả train và test\nBảng 4.7. Kết quả baseline và cấu hình được chọn\nBảng 4.8. Đối chiếu với yêu cầu bài toán"
            paras[i + 1].paragraph_format.first_line_indent = Cm(0)
            paras[i + 1].alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_chapter_1(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN", level=1)
    add_heading(doc, "1.1. Phát biểu bài toán", level=2)
    add_body(doc, "Bài toán được giảng viên giao là xử lý một tập dữ liệu gồm 50 tập tin chứa các bài viết. Mỗi tập tin có thể chứa một hoặc nhiều bài viết, được phân biệt bằng thuộc tính docid; nội dung của từng bài viết được chia thành các câu và lưu trong các thẻ <s>. Yêu cầu đặt ra là đọc toàn bộ dữ liệu, lấy các câu thuộc từng bài viết và trích xuất những câu mang thông tin quan trọng để tạo thành bản tóm tắt ngắn cho bài viết đó.")
    add_body(doc, "Đây là bài toán tóm tắt văn bản theo hướng trích xuất. Hệ thống không sinh câu mới mà lựa chọn nguyên vẹn các câu đã có trong bài viết. Vì vậy, nhiệm vụ chính là xác định mức độ quan trọng của từng câu, chọn các câu có thứ hạng cao và trả chúng về đúng thứ tự xuất hiện trong nguồn. Kết quả thu được được gọi là candidate summary của từng bài viết.")
    add_heading(doc, "1.2. Dữ liệu đầu vào", level=2)
    add_body(doc, "Dữ liệu đầu vào nằm trong thư mục DUC_TEXT/train và gồm 50 tập tin. Notebook vocabularies.ipynb đã đọc toàn bộ 50 tập tin, ghi nhận 14.707 câu, 276.547 token và 18.995 từ khác nhau. Mỗi câu được biểu diễn bởi một thẻ <s> với các trường thông tin chính: docid xác định bài viết, num xác định thứ tự câu, wdcount biểu diễn số từ và content chứa nội dung câu.")
    add_body(doc, "Một tập tin không được xem như một bài viết duy nhất. Sau khi đọc các thẻ <s>, hệ thống phải nhóm các câu theo docid để thu được từng bài viết độc lập. Trong bước thử nghiệm chi tiết, notebook sử dụng tập tin d061j. Tập tin này chứa 186 câu thuộc sáu docid khác nhau; do đó hệ thống tạo sáu bản tóm tắt riêng thay vì gộp toàn bộ 186 câu thành một tài liệu.")
    doc.add_page_break()
    add_table(doc, ["Thành phần dữ liệu", "Ý nghĩa trong bài toán"], [
        ["50 tập tin", "Tập dữ liệu bài viết cần đọc và xử lý"],
        ["docid", "Mã dùng để nhóm các câu thành từng bài viết"],
        ["num", "Thứ tự của câu trong bài viết nguồn"],
        ["wdcount", "Số từ của câu, có thể dùng khi kiểm soát budget"],
        ["content", "Nội dung nguyên bản được giữ lại khi tạo bản tóm tắt"],
    ], [5.0, 10.7])
    add_caption(doc, "Bảng 1.1. Cấu trúc dữ liệu đầu vào của bài toán")
    add_heading(doc, "1.3. Yêu cầu đầu ra", level=2)
    add_body(doc, "Với mỗi docid, đầu ra là một đoạn tóm tắt gồm các câu quan trọng nhất của bài viết. Các câu phải được lấy nguyên văn từ trường content, không được thay đổi nội dung hoặc ghép thành câu mới. Câu được lựa chọn theo điểm xếp hạng, nhưng trước khi nối thành bản tóm tắt phải được sắp xếp lại theo num hoặc source_index để giữ trình tự của bài viết nguồn.")
    add_body(doc, "Notebook hiện tại sử dụng TOP_N = 3, nghĩa là chọn ba câu có điểm PageRank cao nhất cho mỗi bài viết. Khi mở rộng theo budget, hệ thống có thể giới hạn tổng số từ dựa trên wdcount hoặc số token thực tế. Nếu một câu vượt phần budget còn lại, câu đó nên được bỏ qua thay vì cắt ngang, sau đó tiếp tục xét ứng viên kế tiếp.")
    add_heading(doc, "1.4. Phương pháp giải quyết", level=2)
    add_body(doc, "Để giải quyết bài toán, mỗi bài viết được xử lý độc lập theo chuỗi bước sau: lấy và sắp xếp các câu; tiền xử lý nội dung; chuyển mỗi câu thành vector TF–IDF; tính độ tương đồng cosine giữa mọi cặp câu; dùng ma trận tương đồng để tạo đồ thị câu có trọng số; chạy PageRank/TextRank để tính điểm quan trọng; chọn các câu theo score và giới hạn độ dài; cuối cùng khôi phục thứ tự nguồn để tạo candidate summary.")
    add_body(doc, "Trong đồ thị, mỗi câu là một đỉnh và trọng số cạnh là mức tương đồng cosine giữa hai câu. Một câu có liên kết mạnh với nhiều câu quan trọng khác sẽ nhận điểm PageRank cao. Cách tiếp cận này phù hợp với yêu cầu vì không cần nhãn huấn luyện và có thể áp dụng trực tiếp cho từng bài viết trong 50 tập tin.")
    add_heading(doc, "1.5. Mục tiêu thực hiện", level=2)
    add_body(doc, "Báo cáo tập trung vào các mục tiêu cụ thể: (1) đọc đủ 50 tập tin; (2) trích xuất các thẻ câu và nhóm chính xác theo docid; (3) xây dựng biểu diễn TF–IDF cho từng bài viết; (4) tạo ma trận tương đồng và đồ thị có trọng số; (5) xếp hạng câu bằng PageRank; (6) lấy các câu quan trọng làm đầu ra; và (7) phân tích các trường hợp chọn đúng trọng tâm, trùng lặp hoặc mất ngữ cảnh.")
    add_heading(doc, "1.6. Phạm vi báo cáo", level=2)
    add_body(doc, "Phạm vi của báo cáo được giới hạn ở bài toán giảng viên cung cấp và dữ liệu có trong thư mục DUC_TEXT/train. Báo cáo trình bày thống kê của toàn bộ 50 tập tin và phân tích chi tiết kết quả đã lưu đối với d061j. Phương pháp sử dụng TF–IDF, cosine similarity, đồ thị vô hướng có trọng số và PageRank của NetworkX. Báo cáo không mở rộng sang hệ thống ứng dụng khác, không triển khai mô hình sinh và không khẳng định độ chính xác ROUGE khi notebook chưa lưu phép đo này.")
    add_heading(doc, "1.7. Bố cục báo cáo", level=2)
    add_body(doc, "Chương 2 trình bày cơ sở lý thuyết phục vụ trực tiếp cho bài toán: tóm tắt trích xuất, TF–IDF, cosine similarity, đồ thị và PageRank/TextRank. Chương 3 mô tả từng bước xử lý 50 tập tin và thuật toán chọn câu. Chương 4 trình bày môi trường cài đặt, thống kê dữ liệu, kết quả thử nghiệm trên d061j và phân tích lỗi. Chương 5 tổng kết những nội dung đã thực hiện, hạn chế và hướng cải tiến.")


def add_chapter_2(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", level=1)
    add_heading(doc, "2.1. Tóm tắt văn bản trích xuất", level=2)
    add_body(doc, "Tóm tắt trích xuất xem mỗi câu nguồn là một ứng viên. Bài toán có thể được mô hình hóa thành việc ước lượng hàm score(s_i) cho từng câu s_i, sau đó chọn một tập con thỏa ràng buộc độ dài. Vì câu được giữ nguyên, độ trung thực theo nghĩa từ vựng cao: mọi thông tin trong bản tóm tắt đều xuất hiện trong nguồn. Tuy nhiên, việc ghép các câu lấy từ vị trí khác nhau có thể gây đứt mạch tham chiếu, lặp ý hoặc thiếu câu nền.")
    add_body(doc, "Khác với các phương pháp học có giám sát, TextRank khai thác cấu trúc liên kết giữa đơn vị văn bản. Hai câu chia sẻ nhiều từ quan trọng sẽ có độ tương đồng cao và tạo một cạnh mạnh. Một câu được nhiều câu quan trọng khác liên kết tới sẽ nhận điểm cao, tương tự trực giác của PageRank trên Web.")
    add_heading(doc, "2.2. Tiền xử lý và tách câu", level=2)
    add_body(doc, "Trong tập dữ liệu sử dụng, các câu đã được cung cấp bởi thẻ <s>; vì vậy notebook thực hiện trích xuất câu bằng BeautifulSoup thay vì phải tự nhận biết dấu kết câu. Mỗi bản ghi giữ docid, num, wdcount và content. Trường num cần được chuyển sang số nguyên hoặc bảo toàn source_index để tránh lỗi sắp xếp chuỗi, ví dụ “10” đứng trước “9”.")
    add_body(doc, "Tiền xử lý thường gồm chuyển chữ thường, mở rộng dạng viết tắt, tách từ, loại dấu câu và chuẩn hóa khoảng trắng. Stopword có thể được loại bỏ để giảm cạnh được tạo bởi từ chức năng, nhưng cũng cần thử nghiệm vì việc loại bỏ quá mạnh có thể làm mất tín hiệu. Phiên bản notebook hiện tại sử dụng TfidfVectorizer trên chuỗi token đã tiền xử lý; tuy nhiên cell tạo trường preprocessed không còn được lưu, nên đây là điểm tái lập phải sửa trong phiên bản tiếp theo.")
    add_heading(doc, "2.3. Biểu diễn TF–IDF", level=2)
    add_body(doc, "TF–IDF gán trọng số lớn cho một thuật ngữ xuất hiện nhiều trong một câu nhưng ít phổ biến trong tập câu. Với thuật ngữ t và câu d, TF(t,d) đo tần suất của t trong d; IDF(t) giảm trọng số của từ xuất hiện trong nhiều câu. Dạng minh họa trong notebook là:")
    p = doc.add_paragraph(style="Code Academic")
    p.add_run("TF(t,d) = count(t,d) / |d|\nIDF(t) = log(N / df(t))\nTF–IDF(t,d) = TF(t,d) × IDF(t)")
    add_body(doc, "Trong phần cài đặt chính, scikit-learn tạo ma trận thưa X kích thước m × v, với m là số câu hợp lệ và v là số đặc trưng từ vựng. Mỗi hàng là vector của một câu. Khi chuẩn hóa L2, tích vô hướng giữa hai hàng chính là cosine similarity. Cách biểu diễn này dễ giải thích và đủ nhẹ cho thực nghiệm quy mô notebook.")
    add_table(doc,
        ["Đặc trưng", "d1", "d2", "d3", "Ý nghĩa"],
        [
            ["hutech", "0", "0", "0", "Xuất hiện trong mọi tài liệu nên IDF = 0 ở công thức minh họa"],
            ["irs", "0,2197", "0", "0", "Đặc trưng riêng của d1"],
            ["caohoc", "0,0811", "0,0811", "0", "Chung giữa d1 và d2"],
            ["ai", "0", "0,2197", "0", "Đặc trưng riêng của d2"],
            ["oop", "0", "0", "0,2197", "Đặc trưng riêng của d3"],
            ["daihoc", "0", "0", "0,2197", "Đặc trưng riêng của d3"],
        ], [2.2, 1.6, 1.6, 1.6, 8.3])
    add_caption(doc, "Bảng 2.1. Sáu đặc trưng TF–IDF được tính trong ví dụ notebook")
    add_heading(doc, "2.4. Độ tương đồng cosine", level=2)
    add_body(doc, "Độ tương đồng cosine giữa hai vector x và y được tính bởi cos(x,y) = (x·y)/(||x||₂||y||₂). Giá trị gần 1 cho biết hai câu có hướng biểu diễn gần nhau; giá trị 0 cho biết hai câu không chia sẻ đặc trưng sau tiền xử lý. Notebook dùng cosine_similarity để tính đồng thời mọi cặp câu và đặt đường chéo của ma trận bằng 0 nhằm loại cạnh tự nối.")
    add_body(doc, "Ma trận tương đồng S có kích thước m × m và đối xứng khi dùng cosine. Phần tử S_ij trở thành trọng số cạnh giữa câu i và câu j. Notebook giữ cả các giá trị nhỏ thay vì đặt ngưỡng; lựa chọn này giúp đồ thị ít bị rời rạc nhưng có thể tạo nhiều cạnh yếu và tăng nhiễu.")
    add_heading(doc, "2.5. Đồ thị câu có trọng số", level=2)
    add_body(doc, "Đồ thị G = (V,E,W) gồm tập đỉnh V tương ứng với các câu hợp lệ. Một cạnh (i,j) biểu diễn quan hệ tương đồng giữa hai câu và có trọng số w_ij = S_ij. NetworkX xây dựng đồ thị vô hướng từ ma trận NumPy. Trong đồ thị này, câu không được đánh giá riêng lẻ; điểm quan trọng phụ thuộc cả cường độ kết nối và điểm của các láng giềng.")
    add_heading(doc, "2.6. PageRank và TextRank", level=2)
    add_body(doc, "PageRank ban đầu được đề xuất để xếp hạng trang Web dựa trên cấu trúc liên kết. TextRank chuyển trực giác đó sang văn bản: đơn vị văn bản là đỉnh và quan hệ tương đồng là cạnh. Với đồ thị có trọng số, điểm của đỉnh i được cập nhật theo tổng đóng góp từ các đỉnh j nối tới i, có chuẩn hóa theo tổng trọng số cạnh đi ra của j:")
    p = doc.add_paragraph(style="Code Academic")
    p.add_run("PR(i) = (1 − d)/N + d × Σ[j∈N(i)] (w_ji / Σ[k∈N(j)] w_jk) × PR(j)")
    add_body(doc, "Trong đó d là hệ số damping, N là số đỉnh và N(i) là tập láng giềng của i. Notebook sử dụng alpha = 0,85 và thuộc tính weight. Thành phần (1−d)/N bảo đảm xác suất nhảy ngẫu nhiên, hỗ trợ hội tụ và tránh toàn bộ điểm bị giữ trong một vùng đồ thị. Sau hội tụ, các câu được sắp xếp theo score giảm dần.")
    add_heading(doc, "2.7. Chọn câu theo score và budget", level=2)
    add_body(doc, "Xếp hạng chỉ tạo ra thứ tự ưu tiên; hệ thống còn phải quyết định độ dài bản tóm tắt. Cấu hình đơn giản là chọn Top-N. Cấu hình thực tế hơn dùng max_words hoặc số byte, lần lượt duyệt câu theo score và chỉ nhận câu nếu tổng độ dài không vượt budget. Không nên cắt ngang một câu vì sẽ phá vỡ cú pháp; câu quá dài có thể được bỏ qua để xét ứng viên kế tiếp.")
    add_body(doc, "Sau khi chọn, các câu phải được sắp xếp theo source_index trước khi nối lại. Bước này tách hai mục tiêu: score quyết định câu nào đáng giữ, còn thứ tự nguồn quyết định cách trình bày. Đây là một heuristic đơn giản để tăng tính mạch lạc mà không cần mô hình sinh.")


def add_chapter_3(doc: Document, figs: dict[str, Path]) -> None:
    add_heading(doc, "CHƯƠNG 3. GIẢI PHÁP ĐỀ XUẤT", level=1)
    add_heading(doc, "3.1. Kiến trúc tổng thể", level=2)
    add_body(doc, "Giải pháp được tổ chức thành một pipeline tuyến tính: document → tách câu → vector hóa câu → ma trận tương đồng → đồ thị câu có trọng số → PageRank/TextRank → chọn câu theo score và budget → trả câu về thứ tự nguồn → candidate summary. Mỗi bước có đầu vào và đầu ra rõ ràng, nhờ đó có thể kiểm tra riêng lỗi đọc dữ liệu, lỗi biểu diễn, lỗi đồ thị hoặc lỗi lựa chọn.")
    doc.add_picture(str(figs["pipeline"]), width=Cm(16.0))
    add_caption(doc, "Hình 3.1. Luồng xử lý của phương pháp đề xuất")
    add_heading(doc, "3.2. Bước 1 – Đọc và nhóm dữ liệu", level=2)
    add_body(doc, "Notebook đọc file d061j bằng UTF-8, dùng BeautifulSoup tìm toàn bộ thẻ <s> và chuyển từng thẻ thành dictionary. Các câu được nhóm theo docid để mỗi bài báo được tóm tắt độc lập. Trong file d061j có 186 câu thuộc sáu document. Thứ tự câu được lấy từ thuộc tính num; về mặt triển khai cần ép kiểu số hoặc giữ chỉ số xuất hiện ban đầu.")
    add_heading(doc, "3.3. Bước 2 – Tiền xử lý câu", level=2)
    add_body(doc, "Với mỗi câu, quy trình tạo danh sách token: mở rộng contraction, chuyển về chữ thường, chuẩn hóa khoảng trắng, tách từ và giữ token khớp mẫu từ vựng. Câu không có token hợp lệ bị loại khỏi bước vector hóa. Cần lưu cả content gốc để tạo đầu ra và preprocessed để tính toán, tránh làm biến đổi văn bản trả về.")
    add_heading(doc, "3.4. Bước 3 – Vector hóa bằng TF–IDF", level=2)
    add_body(doc, "Tập corpus của một document được tạo bằng cách nối token của từng câu thành chuỗi. TfidfVectorizer.fit_transform(corpus) sinh ma trận TF–IDF. Việc fit riêng theo document làm IDF phản ánh mức độ phân biệt giữa các câu trong chính bài báo, phù hợp mục tiêu chọn câu nội bộ; tuy nhiên trọng số không thể so sánh trực tiếp giữa hai document khác nhau.")
    add_heading(doc, "3.5. Bước 4 – Ma trận tương đồng", level=2)
    add_body(doc, "Hàm cosine_similarity nhận ma trận TF–IDF và trả về S. Đường chéo được đặt bằng 0 vì một câu luôn giống chính nó và cạnh tự nối không mang thông tin xếp hạng. Ví dụ ở Hình 3.2 minh họa năm câu: S1 và S2 có độ tương đồng 0,62 nên tạo cạnh mạnh; S1 và S5 bằng 0 nên không có liên kết có ý nghĩa.")
    doc.add_picture(str(figs["matrix"]), width=Cm(10.5))
    add_caption(doc, "Hình 3.2. Minh họa ma trận tương đồng; số liệu dùng để giải thích, không phải kết quả d061j")
    add_heading(doc, "3.6. Bước 5 – Xây dựng đồ thị", level=2)
    add_body(doc, "Đồ thị được tạo bằng nx.from_numpy_array(similarity_matrix). Mỗi chỉ số dòng trở thành một đỉnh và mọi phần tử khác 0 trở thành cạnh có trọng số. Hình 3.3 cho thấy kích thước đỉnh có thể biểu diễn PageRank và độ dày cạnh biểu diễn cosine similarity. Câu S2 có nhiều liên kết mạnh nên nhận score cao hơn, trong khi S5 nằm ở vùng ngoại biên.")
    doc.add_picture(str(figs["graph"]), width=Cm(11.5))
    add_caption(doc, "Hình 3.3. Minh họa đồ thị câu có trọng số và điểm PageRank")
    add_heading(doc, "3.7. Bước 6 – Xếp hạng bằng PageRank", level=2)
    add_body(doc, "Notebook gọi nx.pagerank(graph, alpha=0.85, weight='weight'). Kết quả là ánh xạ từ chỉ số câu sang score. Danh sách ranked được sắp xếp giảm dần theo score. Nếu document có ít hơn hai câu hợp lệ, hệ thống trả câu duy nhất hoặc chuỗi rỗng; nhánh này giúp tránh tạo đồ thị không cần thiết.")
    add_heading(doc, "3.8. Bước 7 – Lựa chọn và khôi phục thứ tự", level=2)
    add_body(doc, "Ở cấu hình đã chạy, ba chỉ số đầu của ranked được chọn. Sau đó các chỉ số được sắp xếp tăng dần trước khi truy xuất content. Cách làm bảo đảm bản tóm tắt không trình bày theo thứ tự score—một thứ tự có thể đảo diễn biến thời gian—mà theo trật tự kể chuyện của nguồn. Khi bổ sung max_words, bộ chọn cần duyệt theo score, kiểm tra độ dài nguyên câu, bỏ qua ứng viên vượt phần budget còn lại và cuối cùng vẫn sắp xếp theo source_index.")
    add_heading(doc, "3.9. Giả mã thuật toán", level=2)
    pseudo = [
        "INPUT: document D, top_n, max_words (tùy chọn)",
        "sentences ← split_or_parse(D)",
        "valid ← preprocess(sentences) và loại câu rỗng",
        "X ← TFIDF(valid)",
        "S ← cosine_similarity(X); diagonal(S) ← 0",
        "G ← weighted_undirected_graph(S)",
        "scores ← PageRank(G, damping=0.85)",
        "ranked ← sort(valid, key=(−score, source_index))",
        "selected ← chọn nguyên câu theo top_n và/hoặc max_words",
        "selected ← sort(selected, key=source_index)",
        "OUTPUT: join(original_content(selected))",
    ]
    for line in pseudo:
        doc.add_paragraph(line, style="Code Academic")
    add_heading(doc, "3.10. Độ phức tạp và yêu cầu tài nguyên", level=2)
    add_body(doc, "Với m câu và v đặc trưng, TF–IDF thưa phụ thuộc số phần tử khác 0. Việc tạo đầy đủ ma trận tương đồng có chi phí thời gian và bộ nhớ xấp xỉ O(m²), trở thành nút thắt khi document rất dài. PageRank lặp trên đồ thị có chi phí theo số cạnh mỗi vòng. Đối với tin tức vài chục câu như d061j, chi phí này nhỏ; với báo cáo hàng nghìn câu nên đặt ngưỡng similarity hoặc lấy k láng giềng gần nhất để tạo đồ thị thưa.")
    add_heading(doc, "3.11. Các quyết định thiết kế", level=2)
    add_bullets(doc, [
        "Đồ thị vô hướng vì cosine similarity đối xứng.",
        "Giữ trọng số cạnh để PageRank phân biệt liên kết mạnh và yếu.",
        "Đặt đường chéo bằng 0 để loại self-loop.",
        "Tách content gốc khỏi dữ liệu xử lý để không làm biến dạng đầu ra.",
        "Chọn theo score nhưng trình bày theo thứ tự nguồn.",
        "Không cắt ngang câu khi áp dụng word budget."
    ])


def add_chapter_4(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 4. CÀI ĐẶT, THỰC NGHIỆM VÀ ĐÁNH GIÁ", level=1)
    add_heading(doc, "4.1. Môi trường và thư viện", level=2)
    add_body(doc, "Notebook được viết bằng Python và sử dụng pathlib để quản lý đường dẫn, BeautifulSoup để phân tích dữ liệu có thẻ, NumPy cho ma trận, scikit-learn cho TF–IDF/cosine similarity và NetworkX cho đồ thị/PageRank. Đây là các thư viện phổ biến, giúp tập trung vào logic thuật toán thay vì tự cài đặt toàn bộ đại số tuyến tính và vòng lặp hội tụ.")
    add_table(doc, ["Thành phần", "Vai trò"], [
        ["pathlib", "Đọc file theo đường dẫn tương đối và mã hóa UTF-8"],
        ["BeautifulSoup", "Trích xuất thẻ <s> và thuộc tính câu"],
        ["TfidfVectorizer", "Xây dựng ma trận đặc trưng TF–IDF"],
        ["cosine_similarity", "Tạo ma trận tương đồng câu–câu"],
        ["NumPy", "Đặt đường chéo ma trận bằng 0 và xử lý mảng"],
        ["NetworkX", "Tạo đồ thị có trọng số và tính PageRank"],
    ], [4.2, 11.5])
    add_heading(doc, "4.2. Dữ liệu thực nghiệm", level=2)
    add_body(doc, "Notebook vocabularies.ipynb duyệt 50 file trong DUC_TEXT/train và ghi nhận 14.707 câu, 276.547 token, 18.995 từ khác nhau. Hai mươi từ phổ biến nhất chủ yếu là stopword như the, of, and, to, in; quan sát này cho thấy bước IDF và/hoặc loại stopword cần thiết để giảm tác động của từ chức năng.")
    add_table(doc, ["Chỉ tiêu", "Giá trị đã lưu"], [
        ["Số file train", "50"],
        ["Tổng số câu/đoạn được đánh dấu", "14.707"],
        ["Tổng số token", "276.547"],
        ["Số từ khác nhau", "18.995"],
        ["File dùng cho demo tóm tắt", "d061j"],
        ["Số document trong d061j", "6"],
        ["Số câu trong d061j", "186"],
    ], [8.0, 7.7])
    add_caption(doc, "Bảng 4.1. Thống kê dữ liệu được lưu trong hai notebook")
    add_heading(doc, "4.3. Kích thước ma trận theo document", level=2)
    add_body(doc, "Kết quả cell cosine cho biết số câu hợp lệ, số đặc trưng TF–IDF và kích thước ma trận tương đồng của từng document. Tổng số câu hợp lệ là 186, bằng số thẻ câu đã đọc, cho thấy trong phiên chạy tạo output không có câu bị loại vì token rỗng.")
    dims = [
        ["AP880911-0016", "16", "16 × 168", "16 × 16"],
        ["AP880912-0095", "32", "32 × 376", "32 × 32"],
        ["AP880912-0137", "30", "30 × 322", "30 × 30"],
        ["AP880915-0003", "56", "56 × 467", "56 × 56"],
        ["AP880916-0060", "35", "35 × 266", "35 × 35"],
        ["WSJ880912-0064", "17", "17 × 176", "17 × 17"],
    ]
    add_table(doc, ["Document", "Câu hợp lệ", "TF–IDF", "Similarity"], dims, [5.2, 3.0, 3.8, 3.7])
    add_caption(doc, "Bảng 4.2. Kích thước biểu diễn được in bởi executive-summary.ipynb")
    add_heading(doc, "4.4. Kết quả Top-3", level=2)
    add_body(doc, "Mỗi document được tóm tắt bằng ba câu. Thay vì chép toàn bộ sáu bản tóm tắt dài vào phần thân, bảng dưới ghi nhận nội dung trọng tâm và lỗi quan sát được. Các trích đoạn đầy đủ vẫn nằm trong output của notebook, là bằng chứng thực nghiệm gốc.")
    results = [
        ["AP880911-0016", "Hướng di chuyển và tốc độ bão Gilbert; cảnh báo thời tiết; thêm một câu về Florence.", "Câu về Florence lệch chủ đề chính Gilbert, biểu hiện nhiễu do từ vựng thời tiết tương đồng."],
        ["AP880912-0095", "Cảnh báo Tây Caribe; bão mạnh lên; mưa, mất điện và lũ ở Dominican Republic/Haiti.", "Bao phủ diễn biến chính khá tốt nhưng có lặp thông tin bão mạnh lên và mưa lớn."],
        ["AP880912-0137", "Thiệt hại vùng bờ; biện pháp của Hải quân; tình hình Guantanamo.", "Hai câu cuối gần nhau về cùng một cơ sở Hải quân, làm giảm đa dạng nội dung."],
        ["AP880915-0003", "Tàu mắc cạn ở Cancun; vị trí bão; vùng hurricane watch tại Texas/Mexico.", "Ba câu cung cấp sự kiện, vị trí và cảnh báo; độ mạch lạc tương đối tốt."],
        ["AP880916-0060", "Ba phát biểu giải thích điều kiện hình thành bão và năng lượng đại dương.", "Phụ thuộc ngữ cảnh trích dẫn; đại từ và người nói không được giới thiệu đầy đủ."],
        ["WSJ880912-0064", "Vị trí bão; thiệt hại ở Puerto Rico; dự báo bão mạnh và mưa ở Hispaniola.", "Bao phủ tốt nhưng giữ nhiều chi tiết địa danh, chưa tối ưu cho budget ngắn."],
    ]
    add_table(doc, ["Document", "Nội dung được chọn", "Nhận xét/error analysis"], results, [3.3, 6.2, 6.2])
    add_caption(doc, "Bảng 4.3. Phân tích định tính sáu candidate summary Top-3")
    add_heading(doc, "4.5. Phân tích một trường hợp", level=2)
    add_body(doc, "Ở AP880911-0016, hai câu đầu của bản tóm tắt nói trực tiếp về Gilbert: tốc độ gió và hướng di chuyển. Câu thứ ba lại nói Hurricane Florence bị hạ cấp. Câu này có thể đạt PageRank cao vì chia sẻ các từ hurricane, storm, winds và rain với nhiều câu khác. Tuy nhiên về mặt chủ đề, Florence là một sự kiện phụ. Đây là ví dụ điển hình cho giới hạn của TF–IDF: tương đồng từ vựng không đồng nghĩa với cùng thực thể hoặc cùng trọng tâm.")
    add_body(doc, "Ở AP880916-0060, các câu được chọn chứa trích dẫn trực tiếp và các biểu thức “If that happens” hoặc “The sun puts energy…”. Khi tách khỏi câu giới thiệu người nói và bối cảnh, bản tóm tắt khó đọc. PageRank đánh giá độ trung tâm nội dung nhưng không kiểm tra tính độc lập diễn ngôn của từng câu. Một bộ lọc có thể hạ điểm câu bắt đầu bằng đại từ chỉ xuất hoặc một bước mở rộng ngữ cảnh có thể chọn thêm câu đứng trước.")
    add_heading(doc, "4.6. Đánh giá theo tiêu chí", level=2)
    add_body(doc, "Kết quả cho thấy pipeline đã hoàn thành đường đi kỹ thuật từ dữ liệu đến candidate summary. Sáu document đều tạo được ma trận, đồ thị, điểm xếp hạng và bản tóm tắt. Tuy nhiên, “chạy được” không đồng nghĩa với “đạt độ chính xác cao”. Chưa có ROUGE hoặc đánh giá con người được lưu, vì vậy báo cáo chỉ kết luận về tính vận hành và đặc điểm định tính.")
    rubric = [
        ["Mục tiêu và input/output", "Chương 1 xác định document đầu vào, Top-N/budget và candidate summary đầu ra."],
        ["Ý tưởng phương pháp", "Chương 2–3 giải thích tóm tắt trích xuất và trực giác đồ thị."],
        ["Các bước thực hiện", "Mục 3.2–3.9 mô tả lần lượt từng bước và giả mã."],
        ["> 5 đặc trưng / đồ thị", "Bảng 2.1 trình bày 6 đặc trưng; Hình 3.2–3.3 minh họa ma trận và đồ thị."],
        ["Phân lớp/phân nhóm", "Các câu được xếp hạng thành nhóm ưu tiên theo PageRank; Top-N là lớp câu được chọn."],
        ["Lấy đầu vào và tách văn bản", "BeautifulSoup lấy 186 thẻ <s>, nhóm theo docid."],
        ["Nhận xét kết quả", "Mục 4.4–4.6 phân tích ưu điểm, lỗi và giới hạn chứng cứ."],
        ["Cải tiến", "Mục 5.3 đề xuất stopword, threshold, MMR, ROUGE và embedding ngữ nghĩa."],
    ]
    add_table(doc, ["Tiêu chí", "Bằng chứng trong báo cáo"], rubric, [5.0, 10.7])
    add_caption(doc, "Bảng 4.4. Đối chiếu nội dung với rubric trong đề bài")
    add_heading(doc, "4.7. Ưu điểm của phương pháp", level=2)
    add_bullets(doc, [
        "Không cần dữ liệu gán nhãn hoặc quá trình huấn luyện.",
        "Đầu ra truy vết được vì dùng nguyên câu nguồn.",
        "Mỗi thành phần có ý nghĩa rõ ràng và dễ minh họa.",
        "Dùng quan hệ toàn cục giữa các câu thay vì chỉ dựa vào vị trí.",
        "Có thể kiểm soát độ dài bằng Top-N hoặc word budget.",
        "Thích hợp làm baseline cho các mô hình phức tạp hơn."
    ])
    add_heading(doc, "4.8. Hạn chế và đe dọa đối với tính hợp lệ", level=2)
    add_bullets(doc, [
        "TF–IDF không nhận biết đồng nghĩa, thực thể và quan hệ ngữ nghĩa sâu.",
        "PageRank ưu tiên tính trung tâm nên có thể chọn nhiều câu gần nghĩa, gây trùng lặp.",
        "Câu trích dẫn hoặc đại từ có thể mất ngữ cảnh khi đứng độc lập.",
        "Top-3 không thích nghi theo độ dài document hoặc độ dài câu.",
        "Không có ngưỡng similarity nên đồ thị có thể chứa nhiều cạnh yếu.",
        "Output hiện tại chưa kèm score, source_index và số từ để kiểm toán quyết định.",
        "Cell tạo preprocessed không còn trong notebook hiện tại; output đã lưu không đủ chứng minh khả năng chạy lại từ kernel sạch.",
        "Chưa có ROUGE hoặc đánh giá con người, vì vậy chưa thể kết luận độ chính xác định lượng."
    ])
    add_heading(doc, "4.9. Kế hoạch đánh giá định lượng", level=2)
    add_body(doc, "DUC cung cấp summary tham chiếu và các ràng buộc độ dài. Phiên bản tiếp theo nên ghép đúng document với bản tóm tắt người viết, áp dụng cùng budget và tính ROUGE-1, ROUGE-2, ROUGE-L. Cần báo cáo trung bình trên toàn tập cùng độ lệch chuẩn hoặc khoảng tin cậy, đồng thời so sánh ít nhất với baseline Lead-3. Đánh giá con người nên bổ sung các tiêu chí: mức độ đầy đủ, không dư thừa, mạch lạc và khả năng đọc độc lập.")


def add_chapter_5(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    add_heading(doc, "5.1. Kết quả đạt được", level=2)
    add_body(doc, "Đề tài đã mô hình hóa bài toán tóm tắt văn bản trích xuất thành bài toán xếp hạng đỉnh trên đồ thị câu. Pipeline bao quát từ đọc dữ liệu DUC, tách và nhóm câu, biểu diễn TF–IDF, tính cosine similarity, tạo đồ thị có trọng số, chạy PageRank, chọn Top-N và khôi phục thứ tự nguồn. Notebook đã lưu kết quả cho sáu document của d061j với tổng 186 câu và tạo candidate summary ba câu cho mỗi document.")
    add_body(doc, "Về mặt học thuật, báo cáo đã liên kết rõ vai trò của từng thành phần: TF–IDF tạo vector đặc trưng, cosine tạo quan hệ cặp câu, đồ thị tổ chức các quan hệ, PageRank tổng hợp tín hiệu toàn cục và budget chuyển thứ hạng thành bản tóm tắt. Phương pháp đơn giản, minh bạch và phù hợp làm baseline không giám sát.")
    add_heading(doc, "5.2. Những nội dung chưa hoàn thành", level=2)
    add_body(doc, "Hệ thống chưa có phép đo ROUGE được lưu, chưa so sánh với Lead-3 hoặc mô hình khác và chưa đánh giá bởi người dùng. Cơ chế chọn Top-3 chưa xét độ dài, độ dư thừa hoặc tính phụ thuộc ngữ cảnh. Ngoài ra, notebook hiện tại thiếu cell tạo preprocessed trong luồng thực thi, làm giảm khả năng tái lập từ đầu. Những điểm này giới hạn kết luận ở mức pipeline và phân tích định tính.")
    add_heading(doc, "5.3. Hướng cải tiến", level=2)
    add_bullets(doc, [
        "Khôi phục một hàm preprocess duy nhất, chạy notebook từ kernel sạch và lưu phiên bản thư viện.",
        "Chuyển num/wdcount sang số nguyên, giữ source_index ổn định và xuất cả PageRank score.",
        "Thử loại stopword, stemming/lemmatization và n-gram; đánh giá thay đổi bằng cùng tập tham chiếu.",
        "Đặt ngưỡng cosine hoặc xây dựng k-nearest-neighbor graph để loại cạnh yếu và giảm O(m²).",
        "Dùng Maximal Marginal Relevance hoặc phạt tương đồng giữa các câu đã chọn để giảm lặp ý.",
        "Áp dụng word budget, không cắt câu và so sánh các mức nén khác nhau.",
        "Bổ sung đặc trưng vị trí, độ dài, tiêu đề và thực thể; kết hợp chúng với PageRank có cá nhân hóa.",
        "Thay hoặc kết hợp TF–IDF bằng sentence embedding để nhận biết đồng nghĩa và quan hệ ngữ nghĩa.",
        "Đánh giá ROUGE-1/2/L, Lead-3 baseline và đánh giá con người về đầy đủ, mạch lạc, không dư thừa.",
        "Tách notebook thành các hàm/module nhỏ và thêm kiểm thử cho từng giai đoạn để tăng khả năng tái lập."
    ])
    add_heading(doc, "5.4. Kết luận chung", level=2)
    add_body(doc, "TextRank/PageRank là một lựa chọn phù hợp để minh họa cách biến văn bản thành đồ thị và khai thác cấu trúc liên kết nhằm tìm câu quan trọng. Kết quả hiện có cho thấy phương pháp có thể chọn các câu trung tâm và tạo bản tóm tắt nhanh mà không huấn luyện. Đồng thời, error analysis cho thấy cần bổ sung kiểm soát chủ đề, dư thừa, ngữ cảnh và đánh giá tham chiếu. Đây là nền tảng rõ ràng để phát triển từ một notebook minh họa thành một hệ thống tóm tắt có thể tái lập và đánh giá nghiêm ngặt.")


# Revised report content grounded in the executable TextRank pipeline and the
# parameter-analysis notebooks. These definitions intentionally override the
# earlier draft functions while preserving the HUTECH template and formatting.
def add_chapter_1_v2(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN BÀI TOÁN", level=1)
    add_heading(doc, "1.1. Phát biểu bài toán", level=2)
    add_body(doc, "Bài toán cần giải quyết là tự động trích xuất thông tin quan trọng từ 50 tập tin chứa các cụm bài viết tin tức trong thư mục data/DUC_TEXT/train. Mỗi tập tin được xem là một topic; bên trong topic, nội dung đã được chia thành các câu có thẻ <s>. Một câu mang các thuộc tính docid, num và wdcount cùng phần nội dung nguyên văn. Hệ thống phải đọc toàn bộ topic, đánh giá mức độ quan trọng của từng câu trong quan hệ với các câu còn lại, rồi chọn một tập câu ngắn tạo thành candidate summary cho topic.")
    add_body(doc, "Hướng tiếp cận được chọn là tóm tắt trích xuất không giám sát. Hệ thống không sinh ra từ hoặc câu mới, mà giữ nguyên câu nguồn và chỉ quyết định câu nào được chọn. Bài toán vì vậy gồm hai nhiệm vụ liên quan nhưng khác nhau: xếp hạng mức độ quan trọng của câu và lựa chọn một tập câu vừa có thông tin trung tâm, vừa hạn chế trùng lặp, đồng thời không vượt quá ngân sách 100 từ.")
    add_heading(doc, "1.2. Đầu vào, đầu ra và đơn vị xử lý", level=2)
    add_body(doc, "Đơn vị xử lý của pipeline là topic, không phải từng docid riêng lẻ. Mỗi topic có thể tổng hợp nhiều bài viết về cùng một sự kiện; chính quan hệ giữa các câu thuộc các nguồn khác nhau giúp phát hiện nội dung được nhắc lại hoặc được nhiều nguồn cùng xác nhận. Cách tổ chức này phù hợp với dữ liệu DUC và với notebook hiện tại: mỗi tập tin đầu vào sinh đúng một tập tin tóm tắt UTF-8 có cùng tên.")
    add_body(doc, "Mỗi câu được nhận diện ổn định bằng cặp (docid, num). Thuộc tính docid cho biết bài viết nguồn; num là số thứ tự câu trong bài viết. Pipeline còn tạo source_index theo thứ tự xuất hiện thực tế trong topic để có thể khôi phục trình tự sau khi chọn. tagged_content được giữ nguyên để đầu ra vẫn chứa thẻ <s> và metadata, còn content và tokens phục vụ tính toán.")
    doc.add_page_break()
    add_table(doc, ["Thành phần", "Vai trò trong pipeline"], [
        ["50 topic train", "Dữ liệu dùng khảo sát, chọn tham số và tạo tóm tắt"],
        ["9 topic test", "Dữ liệu chỉ dùng đánh giá sau khi khóa cấu hình"],
        ["docid, num", "Khóa định danh câu khi đối chiếu với tập tham chiếu"],
        ["wdcount/content", "Thông tin độ dài và nội dung nguyên văn của câu"],
        ["Candidate summary", "Các câu được chọn, tổng độ dài không quá 100 từ"],
    ], [4.7, 11.0])
    add_caption(doc, "Bảng 1.1. Cấu trúc dữ liệu đầu vào và đầu ra")
    add_heading(doc, "1.3. Mục tiêu thực hiện", level=2)
    add_body(doc, "Mục tiêu thứ nhất là xây dựng một pipeline có thể theo dõi được từ dữ liệu thô đến đầu ra: parse câu, chuẩn hóa và tokenize, tính TF–IDF, tạo ma trận cosine similarity, dựng đồ thị câu có trọng số, tính PageRank, chọn câu bằng PageRank kết hợp MMR, kiểm soát ngân sách và khôi phục thứ tự nguồn. Mã xử lý chính được tách thành các module trong package src/nlp_practice. TF–IDF, cosine, đồ thị adjacency và PageRank được viết tay để quan sát từng phép tính; scikit-learn chỉ được dùng như một implementation tham chiếu để kiểm tra TF–IDF, còn NetworkX chỉ phục vụ trực quan hóa đồ thị.")
    add_body(doc, "Mục tiêu thứ hai là đánh giá và lựa chọn tham số bằng dữ liệu. Bộ notebook thực nghiệm khảo sát ngưỡng similarity, ba tham số hội tụ PageRank và hệ số cân bằng MMR. Cấu hình chỉ được chọn trên tập train. Sau khi khóa cấu hình, hệ thống mới chạy một lần trên tập test để hạn chế việc điều chỉnh theo kết quả kiểm tra.")
    add_heading(doc, "1.4. Câu hỏi cần trả lời", level=2)
    add_bullets(doc, [
        "Có thể biểu diễn mỗi câu bằng TF–IDF và chuyển quan hệ từ vựng giữa các câu thành đồ thị như thế nào?",
        "Ngưỡng cosine similarity tác động ra sao đến mật độ, nút cô lập, số thành phần liên thông và chất lượng câu được chọn?",
        "Damping, tolerance và max iterations ảnh hưởng như thế nào đến thứ hạng và khả năng hội tụ PageRank?",
        "MMR có giúp cân bằng giữa độ quan trọng và độ không trùng lặp của các câu hay không?",
        "Cấu hình được chọn trên train có duy trì chất lượng trên 9 topic test hay không?",
    ])
    add_heading(doc, "1.5. Phạm vi và giới hạn kết luận", level=2)
    add_body(doc, "Báo cáo chỉ sử dụng dữ liệu, cấu hình, chỉ số và output đã xuất hiện trong các notebook pipeline, evaluation và textrank-parameter-analysis. Các số liệu từ notebook cũ không còn phản ánh pipeline hiện tại, các ví dụ số tự tạo và các phép đo chưa được chạy đều bị loại bỏ. Đánh giá chính dựa trên việc một câu dự đoán có trùng cặp (docid, num) với tập câu tham chiếu của cùng topic hay không; không sử dụng ROUGE.")
    add_body(doc, "Đối sánh chính xác ở mức định danh câu có ưu điểm là rõ ràng và kiểm tra được, nhưng không đo được câu diễn đạt tương đương, mức độ đúng về ngữ nghĩa hoặc tính mạch lạc của toàn bản tóm tắt. Vì thế, Precision, Recall và F1 trong báo cáo phải được hiểu là chất lượng exact-sentence extraction trong phạm vi ngân sách 100 từ, không phải độ chính xác ngữ nghĩa tuyệt đối.")
    add_heading(doc, "1.6. Bố cục báo cáo", level=2)
    add_body(doc, "Chương 2 trình bày chi tiết nền tảng lý thuyết của tóm tắt trích xuất, TF–IDF, cosine similarity, đồ thị, PageRank, MMR và các metric. Chương 3 mô tả thiết kế pipeline và quy trình chọn tham số. Chương 4 trình bày thực nghiệm, kết quả so sánh train/test và phân tích lỗi. Chương 5 tổng kết thành tựu, hạn chế và hướng phát triển tiếp theo.")


def add_chapter_2_v2(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT", level=1)
    add_heading(doc, "2.1. Tóm tắt văn bản trích xuất", level=2)
    add_body(doc, "Tóm tắt trích xuất xem tập câu nguồn S = {s₁, s₂, …, sₙ} là không gian ứng viên. Hệ thống ước lượng độ quan trọng r(sᵢ) cho từng câu rồi tìm tập con A ⊆ S thỏa ràng buộc độ dài. Vì câu được giữ nguyên, đầu ra có khả năng truy vết cao: có thể xác định chính xác câu đến từ docid nào và vị trí num nào. Đây là đặc tính phù hợp với yêu cầu trích xuất thông tin từ tập bài viết.")
    add_body(doc, "Tuy nhiên, bài toán không thể giải quyết chỉ bằng cách lấy các câu có score cao nhất. Hai câu cùng mô tả một sự kiện thường cùng có score cao và dễ gây dư thừa. Một câu đứng độc lập có thể chứa đại từ hoặc tham chiếu thiếu ngữ cảnh. Do đó, chất lượng tóm tắt chịu tác động của ít nhất ba yếu tố: tính quan trọng, tính đa dạng và khả năng đọc sau khi ghép câu. Pipeline hiện tại xử lý hai yếu tố đầu bằng PageRank và MMR; thứ tự nguồn được khôi phục như một heuristic hỗ trợ yếu tố thứ ba.")
    add_heading(doc, "2.2. Chuẩn hóa và tokenize", level=2)
    add_body(doc, "Trong mã nguồn hiện tại, mô-đun preprocessing.py dùng BeautifulSoup để đọc các thẻ <s>. Mỗi câu được lưu dưới dạng từ điển gồm content, token và các siêu dữ liệu cần thiết. Nội dung được chuyển về chữ thường bằng lower(), sau đó tokenizer lấy các chuỗi chữ và số theo biểu thức chính quy. Stopword tiếng Anh được loại khỏi danh sách token dùng cho tính toán; trường content vẫn giữ câu nguồn để có thể hoàn nguyên đúng nội dung khi tạo bản tóm tắt.")
    add_body(doc, "Loại stopword làm giảm các quan hệ giả tạo bởi những từ chức năng xuất hiện trong hầu hết câu. Mặt khác, phép xử lý dựa trên từ bề mặt chưa gộp biến thể hình thái, chưa nhận biết đồng nghĩa và chưa giải quyết thực thể đồng tham chiếu. Đây là nguyên nhân nền tảng khiến cosine TF–IDF có thể đánh giá thấp hai câu cùng nghĩa nhưng dùng từ khác, hoặc đánh giá cao hai câu dùng chung từ nhưng nói về sự kiện khác.")
    add_heading(doc, "2.3. Term Frequency và Inverse Document Frequency", level=2)
    add_body(doc, "Trong từng topic, mỗi câu được xem như một document nhỏ. Hàm compute_tf trong src/nlp_practice/tfidf.py dùng tần suất thô, tức số lần thuật ngữ t xuất hiện trong câu d. Cách định nghĩa này trùng với mặc định của TfidfVectorizer khi sublinear_tf=False:")
    p = doc.add_paragraph(style="Code Academic"); p.add_run("TF(t,d) = count(t,d)")
    add_body(doc, "Document Frequency df(t) là số câu chứa t. Hàm compute_df chuyển token của từng câu thành set trước khi cộng, nhờ đó một thuật ngữ chỉ làm df tăng một lần trong mỗi câu. Hàm compute_idf dùng smooth IDF, tương ứng với smooth_idf=True của scikit-learn:")
    p = doc.add_paragraph(style="Code Academic"); p.add_run("IDF(t) = log((1 + N) / (1 + df(t))) + 1")
    add_body(doc, "Trọng số chưa chuẩn hóa là TF–IDF(t,d) = TF(t,d) × IDF(t). Một từ xuất hiện nhiều trong một câu nhưng ít phổ biến ở các câu khác nhận trọng số lớn. Vector được lưu thưa dưới dạng dict[token, weight], vì mỗi câu chỉ chứa một phần nhỏ từ vựng của topic.")
    add_heading(doc, "2.4. Chuẩn hóa L2 và cosine similarity", level=2)
    add_body(doc, "Hàm normalize_l2 tính độ dài Euclid của vector v rồi chia từng trọng số cho độ dài đó. Với v = (w₁,…,wₘ), chuẩn L2 và trọng số sau chuẩn hóa được tính như sau:")
    p = doc.add_paragraph(style="Code Academic"); p.add_run("||v||₂ = sqrt(Σᵢwᵢ²)    và    v̂ᵢ = wᵢ / ||v||₂")
    add_body(doc, "Sau phép chia, vector khác 0 có ||v̂||₂ = 1. Nếu vector rỗng hoặc có độ dài bằng 0, hàm trả lại vector ban đầu để tránh chia cho 0. Chuẩn hóa L2 được bổ sung vì ba nguyên nhân. Thứ nhất, nó loại ảnh hưởng của độ lớn: câu có nhiều token hoặc một từ lặp nhiều không tự động có vector dài hơn và lấn át câu ngắn. Thứ hai, mọi vector nằm trên mặt cầu đơn vị nên có cùng thang đo, thuận tiện cho việc đặt ngưỡng cạnh và đối chiếu implementation. Thứ ba, khi cả hai vector đã có norm bằng 1, cosine similarity rút gọn thành tích vô hướng, giúp tính toán và diễn giải đơn giản hơn.")
    add_body(doc, "Về mặt toán học, L2 không phải điều kiện bắt buộc để hàm cosine đầy đủ trả kết quả đúng, vì similarity.py vẫn chia cho ||x||₂||y||₂. Tuy nhiên, L2 là một quyết định cần thiết trong thiết kế biểu diễn hiện tại: nó làm vector viết tay tương thích trực tiếp với mặc định norm='l2' của TfidfVectorizer và bảo đảm cùng một vector có thể được dùng an toàn cho cả cosine, tích vô hướng và các bước kiểm tra sau này.")
    p = doc.add_paragraph(style="Code Academic"); p.add_run("cos(x,y) = (x · y) / (||x||₂||y||₂); nếu ||x||₂=||y||₂=1 thì cos(x,y)=x·y")
    add_heading(doc, "2.5. Kiểm chứng TF–IDF viết tay bằng scikit-learn", level=2)
    add_body(doc, "Để kiểm tra implementation, main.py thay dữ liệu DUC bằng ba câu nhỏ có kết quả dễ truy vết. Các token được đưa trực tiếp vào cả build_tfidf_vectors và compute_tfidf_with_sklearn. Hàm tham chiếu nối token thành chuỗi, gọi TfidfVectorizer với các mặc định use_idf=True, smooth_idf=True, sublinear_tf=False và norm='l2', rồi chuyển ma trận thưa về cùng dạng dict[token, weight].")
    add_table(doc, ["Câu", "Nội dung", "Token"], [
        ["S1", "cat dog dog", "cat, dog, dog"],
        ["S2", "cat fish", "cat, fish"],
        ["S3", "dog bird", "dog, bird"],
    ], [2.0, 6.0, 7.7])
    add_caption(doc, "Bảng 2.1. Dữ liệu kiểm tra TF–IDF viết tay")
    add_body(doc, "Trong ba câu, df(cat)=df(dog)=2 và df(fish)=df(bird)=1. Do đó IDF của cat và dog bằng 1,287682; IDF của fish và bird bằng 1,693147. Kết quả trước và sau chuẩn hóa cho thấy từ dog xuất hiện hai lần trong S1 nên trọng số thô gấp đôi cat, nhưng hướng của vector được giữ nguyên sau L2.")
    add_table(doc, ["Câu", "TF–IDF chưa L2", "||v||₂", "Vector sau L2"], [
        ["S1", "cat=1,287682; dog=2,575364", "2,879345", "cat=0,447214; dog=0,894427"],
        ["S2", "cat=1,287682; fish=1,693147", "2,127175", "cat=0,605349; fish=0,795961"],
        ["S3", "dog=1,287682; bird=1,693147", "2,127175", "dog=0,605349; bird=0,795961"],
    ], [1.5, 6.0, 2.5, 5.7])
    add_caption(doc, "Bảng 2.2. Kết quả TF–IDF trước và sau chuẩn hóa L2")
    add_table(doc, ["Câu", "Vector viết tay", "Vector scikit-learn", "Sai khác lớn nhất"], [
        ["S1", "cat=0,447214; dog=0,894427", "cat=0,447214; dog=0,894427", "0,0"],
        ["S2", "cat=0,605349; fish=0,795961", "cat=0,605349; fish=0,795961", "0,0"],
        ["S3", "dog=0,605349; bird=0,795961", "dog=0,605349; bird=0,795961", "0,0"],
    ], [1.5, 5.4, 5.4, 3.4])
    add_caption(doc, "Bảng 2.3. Đối chiếu TF–IDF viết tay với scikit-learn")
    add_body(doc, "Hai implementation khớp ở toàn bộ trọng số được kiểm tra; sai khác tuyệt đối lớn nhất bằng 0,0 trong phiên chạy hiện tại. Từ các vector này, cosine(S1,S2)=0,270720; cosine(S1,S3)=0,541440 và cosine(S2,S3)=0. Các giá trị phù hợp trực giác: S1 và S3 chia sẻ dog, trong đó dog có trọng số lớn trong S1; S2 và S3 không chia sẻ token.")
    add_body(doc, "Phép so sánh chỉ chứng minh sự tương đương trên dữ liệu kiểm tra và trong điều kiện cấu hình nêu trên. Với dữ liệu khác, tokenizer mặc định của TfidfVectorizer có thể bỏ token một ký tự hoặc chuẩn hóa khác package tự viết. Vì vậy khi mở rộng kiểm thử cần truyền tokenizer/vocabulary tương thích hoặc tiếp tục so sánh trên chính danh sách token đã kiểm soát.")
    add_heading(doc, "2.6. Đồ thị câu có trọng số", level=2)
    add_body(doc, "Đồ thị G = (V,E,W) là đồ thị vô hướng. Mỗi đỉnh i ∈ V tương ứng một câu. Cạnh (i,j) được tạo khi i ≠ j và cosine(i,j) ≥ τ, với trọng số wᵢⱼ bằng chính cosine similarity. Không có self-loop. Vì cosine đối xứng nên wᵢⱼ = wⱼᵢ. Adjacency list chỉ lưu các cạnh vượt ngưỡng, giúp đồ thị thưa hơn ma trận đầy đủ.")
    add_body(doc, "Bốn chỉ số được dùng để theo dõi sức khỏe đồ thị. Với N đỉnh và E cạnh, density = 2E/[N(N−1)] đo tỷ lệ cạnh được giữ; average degree = 2E/N cho biết số láng giềng trung bình; isolated-node ratio là tỷ lệ đỉnh bậc 0; connected components là số nhóm đỉnh có đường đi liên kết. Threshold thấp làm đồ thị dày và có thể giữ cạnh nhiễu. Threshold cao làm tăng nút cô lập và phân mảnh đồ thị, khiến PageRank ít tín hiệu liên kết để phân biệt câu.")
    add_heading(doc, "2.7. PageRank có trọng số", level=2)
    add_body(doc, "TextRank áp dụng nguyên lý PageRank lên đồ thị văn bản: một câu quan trọng nếu nó được nối mạnh với nhiều câu quan trọng khác. Tại mỗi vòng lặp, điểm của đỉnh i nhận đóng góp từ các láng giềng j theo tỷ lệ giữa trọng số cạnh wⱼᵢ và tổng trọng số đi ra của j:")
    p = doc.add_paragraph(style="Code Academic"); p.add_run("PR(i) = (1−d)/N + d × [dangling/N + Σⱼ wⱼᵢ PR(j) / Σₖwⱼₖ]")
    add_body(doc, "Hệ số damping d điều chỉnh mức phụ thuộc vào đồ thị. Thành phần (1−d)/N đóng vai trò teleportation. Nếu một đỉnh không có cạnh, điểm dangling của nó được gom và phân phối đều cho mọi đỉnh, tránh thất thoát tổng xác suất. Score được khởi tạo đồng đều 1/N; sau mỗi vòng, thuật toán tính tổng sai khác tuyệt đối giữa vector mới và cũ. Quá trình dừng khi sai khác nhỏ hơn tolerance hoặc báo không hội tụ nếu vượt max_iterations.")
    add_body(doc, "Damping lớn làm thứ hạng nhạy hơn với cấu trúc liên kết; damping nhỏ kéo các score gần phân phối đều. Tolerance là tiêu chuẩn hội tụ số học, không phải metric chất lượng. max_iterations là hàng rào an toàn. Do đó, lựa chọn PageRank phải vừa bảo đảm hội tụ trên mọi topic, vừa kiểm tra xem thay đổi sai số có thực sự làm thay đổi thứ hạng câu hay chỉ tăng chi phí tính toán.")
    add_heading(doc, "2.8. Maximal Marginal Relevance", level=2)
    add_body(doc, "MMR chọn câu tuần tự để cân bằng relevance và novelty. Relevance là PageRank sau khi min–max normalization. Với tập câu đã chọn A, redundancy của ứng viên i được lấy bằng similarity lớn nhất giữa i và một câu trong A. Score MMR được tính như sau:")
    p = doc.add_paragraph(style="Code Academic"); p.add_run("MMR(i) = λ × relevance(i) − (1−λ) × max[j∈A] similarity(i,j)")
    add_body(doc, "Khi λ gần 1, hệ thống gần như chỉ ưu tiên PageRank. Khi λ nhỏ, hình phạt trùng lặp mạnh hơn và có thể bỏ qua câu trung tâm. Trường hợp chưa chọn câu nào, phần redundancy bằng 0. Notebook thử MMR tắt và các λ = 0,30; 0,50; 0,70; 0,90; 1,00. Cấu hình λ = 1 tương đương chọn thuần relevance, vì phần phạt bị nhân với 0.")
    add_heading(doc, "2.9. Ràng buộc 100 từ và thứ tự nguồn", level=2)
    add_body(doc, "Bộ chọn duyệt các ứng viên theo PageRank hoặc MMR. Một câu chỉ được nhận nếu thêm toàn bộ câu vẫn không vượt MAX_SUMMARY_WORDS = 100. Nếu câu quá dài so với phần ngân sách còn lại, câu bị bỏ qua và thuật toán tiếp tục xét ứng viên khác; không cắt ngang câu. Sau khi kết thúc, các chỉ số được chọn được sắp theo source_index và tagged_content được ghép lại. Vì vậy thứ tự tính điểm quyết định tập câu, còn thứ tự nguồn quyết định cách trình bày.")
    add_heading(doc, "2.10. Định nghĩa metric đánh giá", level=2)
    add_body(doc, "Với mỗi topic, tập khóa dự đoán P và tập khóa tham chiếu R gồm các cặp (docid, num). Trùng khóa trong reference chỉ tính một lần; dự đoán trùng cũng chỉ giữ lần xuất hiện đầu tiên. True positive là |P ∩ R|. Precision đo tỷ lệ câu dự đoán nằm trong reference; Recall đo tỷ lệ câu reference được thu hồi; F1 là trung bình điều hòa của Precision và Recall.")
    p = doc.add_paragraph(style="Code Academic"); p.add_run("Precision = TP/|P|    Recall = TP/|R|    F1 = 2PR/(P+R)")
    add_body(doc, "Macro metric được tính riêng cho từng topic hợp lệ rồi lấy trung bình, để topic có nhiều câu không lấn át topic nhỏ. Hit@K bằng 1 nếu trong K câu đầu có ít nhất một câu đúng, ngược lại bằng 0; macro Hit@K là tỷ lệ topic có hit. Ngoài metric chất lượng, notebook còn ghi word-budget utilization, redundancy, runtime, số vòng PageRank và các chỉ số đồ thị.")
    add_table(doc, ["Nhóm", "Chỉ số", "Ý nghĩa"], [
        ["Chất lượng", "Macro Precision, Recall, F1", "Đối sánh chính xác câu theo từng topic"],
        ["Khả năng tìm đúng", "Hit@1, Hit@3, Hit@5", "Topic có ít nhất một câu đúng trong K vị trí"],
        ["Tóm tắt", "Word count, budget utilization", "Mức sử dụng giới hạn 100 từ"],
        ["Đa dạng", "Mean/max redundancy", "Tương đồng giữa các câu đã chọn"],
        ["Đồ thị", "Density, isolated ratio, degree, components", "Mức liên kết và phân mảnh của đồ thị"],
        ["Tính toán", "Iterations, runtime, convergence", "Chi phí và độ ổn định thuật toán"],
    ], [3.0, 5.3, 7.4])
    add_caption(doc, "Bảng 2.4. Nhóm chỉ số đánh giá trong notebook")


def add_chapter_3_v2(doc: Document, figs: dict[str, Path]) -> None:
    add_heading(doc, "CHƯƠNG 3. GIẢI PHÁP VÀ PIPELINE THỰC HIỆN", level=1)
    add_heading(doc, "3.1. Kiến trúc tổng thể", level=2)
    add_body(doc, "Pipeline mục tiêu được thiết kế tuyến tính: topic → parse câu → tokenize → sparse TF–IDF → cosine similarity → weighted sentence graph → PageRank → MMR và word budget → sắp theo source_index → candidate summary. Phiên bản hiện tại chuyển phần code thực hành ra package src/nlp_practice để mỗi giai đoạn có một module riêng, còn notebook giữ vai trò mô tả, chạy thí nghiệm và lưu artifact.")
    doc.add_picture(str(figs["pipeline"]), width=Cm(16.0))
    add_caption(doc, "Hình 3.1. Luồng xử lý tóm tắt trích xuất")
    add_table(doc, ["Module", "Trách nhiệm hiện tại"], [
        ["preprocessing.py", "Đọc thẻ <s> bằng BeautifulSoup, tokenize, lọc stopword và giữ metadata"],
        ["tfidf.py", "TF/DF/smooth IDF viết tay, chuẩn hóa L2 và đối chiếu TfidfVectorizer"],
        ["similarity.py", "Tích vô hướng, độ dài vector, cosine và ma trận similarity"],
        ["graph.py", "Tạo adjacency list có trọng số; dùng NetworkX/Matplotlib chỉ để vẽ"],
        ["pagerank.py", "PageRank có dangling redistribution và xếp hạng câu"],
        ["selection.py", "Chuẩn hóa Min–Max điểm PageRank; MMR/budget chưa chuyển đầy đủ"],
        ["main.py", "Điều phối; hiện dừng sau phép so sánh TF–IDF viết tay–sklearn"],
    ], [4.2, 11.5])
    add_caption(doc, "Bảng 3.1. Cấu trúc package src/nlp_practice")
    add_heading(doc, "3.2. Đọc dữ liệu và bảo toàn metadata", level=2)
    add_body(doc, "Hàm read_topic trong preprocessing.py đọc UTF-8 và dùng BeautifulSoup tìm từng thẻ <s>. Mỗi câu được lưu bằng dict gồm docid, num, wdcount, content, tagged_content, token và source_index. num và wdcount được chuyển sang int; source_index nhận vị trí theo thứ tự đọc. Nội dung rỗng bị bỏ qua, còn content gốc được giữ riêng để bước tiền xử lý không làm thay đổi câu xuất ra.")
    add_heading(doc, "3.3. Tạo biểu diễn câu và kiểm chứng", level=2)
    add_body(doc, "tfidf.py dùng các vòng lặp và dict thông thường để người học có thể theo dõi từng bước. compute_tf đếm số lần xuất hiện; compute_df đếm số câu chứa term; compute_idf áp dụng smooth IDF; compute_tfidf nhân TF với IDF; normalize_l2 chuẩn hóa từng vector; build_tfidf_vectors điều phối các hàm. compute_tfidf_with_sklearn là nhánh kiểm chứng độc lập, không thay thế implementation viết tay.")
    add_body(doc, "main.py hiện chọn ba câu cat–dog–fish–bird làm fixture kiểm tra, gọi đồng thời hai implementation và in vector. Lệnh .venv/bin/python -m nlp_practice.main đã xác nhận ba cặp vector giống nhau. Đây là phép kiểm tra cục bộ cho tầng biểu diễn trước khi tiếp tục chạy similarity và đồ thị.")
    add_heading(doc, "3.4. Tính similarity và dựng đồ thị", level=2)
    add_body(doc, "similarity.py tính dot product bằng cách duyệt các term của vector thứ nhất và tra trọng số trong vector thứ hai; vector_magnitude tính chuẩn Euclid; cosine_similarity dùng công thức đầy đủ và trả 0 nếu một vector có độ dài 0. Với n câu, build_similarity_matrix xét từng cặp i<j, đặt giá trị đối xứng và giữ 0 trên đường chéo. graph.py chuyển ma trận thành adjacency list vô hướng khi similarity đạt threshold.")
    add_body(doc, "NetworkX không tham gia tính PageRank của pipeline hiện tại. Nó chỉ được gọi trong draw_sentence_graph để tạo đối tượng nx.Graph, bố trí spring_layout và xuất hình PNG. Sự phân tách này giúp code thuật toán vẫn đọc được bằng các cấu trúc Python cơ bản, đồng thời tận dụng thư viện cho trực quan hóa.")
    add_heading(doc, "3.5. Tính PageRank từ đầu", level=2)
    add_body(doc, "Mỗi vòng PageRank trước hết tính tổng điểm của dangling nodes. Với từng đỉnh, score mới gồm teleportation, phần dangling chia đều và tổng đóng góp có trọng số từ hàng xóm. Tổng outgoing weight được tính trước để tránh lặp. Nếu damping không thuộc (0,1), hàm báo ValueError. Nếu không đạt tolerance trong giới hạn, hàm báo lỗi hội tụ; notebook thực nghiệm giữ trạng thái not_converged thay vì làm mất toàn bộ kết quả sweep.")
    add_heading(doc, "3.6. Chọn câu bằng MMR và budget", level=2)
    add_body(doc, "Notebook đã thực nghiệm PageRank score sau Min–Max normalization, MMR và budget 100 từ. Trong package hiện tại, selection.py mới chuyển normalize_scores; phần vòng lặp MMR, kiểm tra câu có vừa ngân sách và khôi phục source_index chưa được chuyển đầy đủ. Vì vậy các metric toàn DUC trong Chương 4 là bằng chứng từ notebook đã thực thi, còn src/nlp_practice hiện là phiên bản module hóa đang được hoàn thiện.")
    add_heading(doc, "3.7. Cấu hình baseline và cấu hình cuối", level=2)
    add_body(doc, "Pipeline ban đầu dùng threshold 0,10; damping 0,85; tolerance 10⁻⁸; max iterations 1000; MMR bật với λ = 0,70; budget 100 từ. Phép quét thô trước đây chọn 0,02 cho pipeline MMR. Lần cập nhật này bổ sung lưới mịn ở vùng 0–0,03 và khóa threshold 0,0125 theo kết quả PageRank + Top-K=15; damping, tolerance và λ được giữ nguyên để cô lập ảnh hưởng của ngưỡng.")
    add_table(doc, ["Tham số", "Baseline", "Cấu hình cuối", "Vai trò"], [
        ["Similarity threshold", "0,10", "0,0125", "Quyết định cạnh được giữ"],
        ["PageRank damping", "0,85", "0,85", "Mức phụ thuộc vào đồ thị"],
        ["Tolerance", "10⁻⁸", "10⁻⁸", "Điều kiện dừng"],
        ["Max iterations", "1000", "300", "Giới hạn vòng lặp"],
        ["MMR λ", "0,70", "0,70", "Cân bằng relevance–redundancy"],
        ["Word budget", "100", "100", "Độ dài tối đa của summary"],
    ], [4.0, 2.4, 3.0, 6.3])
    add_caption(doc, "Bảng 3.2. Tham số của cấu hình cuối")
    add_heading(doc, "3.8. Quy trình lựa chọn tham số", level=2)
    add_body(doc, "Notebook 01 audit dữ liệu và đo baseline. Notebook 02 quét mười threshold từ 0 đến 0,30. Notebook 03 cố định threshold tốt nhất và thử 5 damping × 3 tolerance × 3 max_iterations. Notebook 04 so sánh MMR tắt với năm λ. Script fine_threshold_sweep.py bổ sung 13 mốc từ 0 đến 0,03, xuất riêng kết quả Top-K=15 và MMR/100 từ để không trộn hai hợp đồng đánh giá.")
    add_body(doc, "Quy tắc xếp hạng là Macro Precision giảm dần, tiếp theo Macro F1 giảm dần, mean redundancy tăng dần và mean runtime tăng dần. Chỉ cấu hình hội tụ trên toàn bộ 50 topic train mới được xét. Cấu hình đứng đầu train được khóa trước khi chạy test. Trình tự này phân tách tuning và final evaluation, làm giảm nguy cơ chọn tham số vì phù hợp ngẫu nhiên với 9 topic test.")
    add_heading(doc, "3.9. Giả mã pipeline", level=2)
    for line in [
        "INPUT: topic T, configuration C, max_words = 100",
        "sentences ← parse_s_tags(T) và giữ (docid, num, source_index)",
        "tokens ← casefold, regex-tokenize, remove_stopwords(sentences)",
        "vectors ← L2_normalize(TF × smooth_IDF(tokens))",
        "similarity ← pairwise_cosine(vectors)",
        "graph ← {(i,j,w) | i ≠ j và similarity[i,j] ≥ threshold}",
        "pagerank ← iterate(graph, damping, tolerance, max_iterations)",
        "selected ← sequential_MMR(pagerank, similarity, max_words)",
        "selected ← sort(selected, key=source_index)",
        "OUTPUT: join(tagged_content[selected])",
    ]:
        doc.add_paragraph(line, style="Code Academic")
    add_heading(doc, "3.10. Tính tái lập và kiểm soát sai lệch", level=2)
    add_body(doc, "Các topic trong notebook được sort; tie-break và quy tắc xếp hạng cấu hình được khai báo; không có lấy mẫu ngẫu nhiên; CSV lưu đầy đủ tham số, metric theo topic và trạng thái hội tụ. Đối với package mới, phép kiểm tra TF–IDF sử dụng fixture cố định và so sánh từng term. Hiện tests/test_main.py chưa có ca kiểm thử tự động, vì vậy kết quả đối chiếu được xác nhận bằng lệnh chạy main chứ chưa được xem là một regression test độc lập.")
    add_body(doc, "Một giới hạn cần nêu rõ là main.py có return ngay sau khi in hai bộ vector. Do đó mã similarity, graph và PageRank nằm phía dưới chưa được thực thi trong chế độ kiểm tra hiện tại. Khi hoàn tất chuyển đổi, nên tách phép kiểm tra TF–IDF thành test hoặc notebook riêng, loại bỏ return tạm thời và để main chạy pipeline DUC từ đầu đến cuối.")
    add_heading(doc, "3.11. Độ phức tạp", level=2)
    add_body(doc, "TF–IDF thưa phụ thuộc tổng số token khác 0. Phần pairwise cosine có độ phức tạp theo số cặp câu, xấp xỉ O(n²), và ma trận similarity cần O(n²) bộ nhớ. PageRank có chi phí O(I×E), với I là số vòng và E là số cạnh vượt threshold. MMR có thể cần so sánh mỗi ứng viên với tập đã chọn; vì summary bị giới hạn 100 từ, số câu đã chọn thường nhỏ. Ở quy mô DUC, thiết kế đủ nhẹ; với topic rất dài cần dùng k-nearest-neighbor graph hoặc không lưu toàn bộ ma trận.")


def add_chapter_4_v2(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 4. THỰC NGHIỆM VÀ ĐÁNH GIÁ", level=1)
    add_heading(doc, "4.1. Thiết kế thực nghiệm", level=2)
    add_body(doc, "Dữ liệu được chia thành 50 topic train và 9 topic test. Mọi lựa chọn tham số dựa trên train. Trong 50 topic train, 34 topic có reference hợp lệ để tính exact-sentence metric; 16 topic có reference rỗng nên vẫn được chạy pipeline và kiểm tra hội tụ nhưng không tham gia macro Precision, Recall, F1. Tập test có 9 topic hợp lệ.")
    add_table(doc, ["Notebook", "Phép thử", "Artifact chính"], [
        ["01 Dataset and baseline", "Audit 50 train/9 test và đo baseline", "Dataset summary, baseline metrics"],
        ["02 Similarity threshold", "10 giá trị từ 0,00 đến 0,30", "Quality, graph health, similarity distribution"],
        ["03 PageRank parameters", "45 tổ hợp damping/tolerance/iterations", "Quality và convergence"],
        ["04 MMR parameters", "MMR off và 5 giá trị λ", "Quality và redundancy"],
        ["05 Final configuration", "24 cấu hình local grid, khóa train rồi chạy test", "Recommended config, train/test metrics"],
    ], [4.2, 5.5, 6.0])
    add_caption(doc, "Bảng 4.1. Thiết kế thực nghiệm theo từng notebook")
    add_heading(doc, "4.2. Kiểm tra TF–IDF viết tay với scikit-learn", level=2)
    add_body(doc, "Phép kiểm tra được chạy bằng Python 3.14.4 và scikit-learn 1.9.0 qua lệnh .venv/bin/python -m nlp_practice.main. Với đúng ba câu ở Bảng 2.1, chương trình in hai danh sách vector: kết quả build_tfidf_vectors viết tay và kết quả compute_tfidf_with_sklearn. Hai danh sách trùng nhau theo từng term; sai khác tuyệt đối lớn nhất là 0,0.")
    add_body(doc, "Kiểm tra này xác nhận bốn quyết định cài đặt đồng nhất với TfidfVectorizer mặc định: dùng TF dạng số đếm, smooth IDF có cộng 1, không dùng sublinear TF và chuẩn hóa L2. Nếu bỏ normalize_l2 ở nhánh viết tay, trọng số còn ở thang chưa chuẩn hóa và không thể so sánh trực tiếp với output mặc định của sklearn. Đây là lý do thực nghiệm cụ thể cho việc bổ sung L2 vào build_tfidf_vectors.")
    add_body(doc, "Kết quả không được khái quát thành việc implementation đúng cho mọi input. Fixture chỉ có bốn token, không có stopword, token một ký tự, Unicode hoặc vector rỗng. Cần bổ sung test cho các trường hợp biên và dùng math.isclose với tolerance rõ ràng thay vì chỉ so sánh output in ra.")
    add_heading(doc, "4.3. Baseline", level=2)
    add_body(doc, "Baseline với threshold 0,10 đạt Macro Precision 0,0918, Macro Recall 0,0383 và Macro F1 0,0534 trên 34 topic train hợp lệ. Hit@1 bằng 0,2353; Hit@3 bằng 0,3824; Hit@5 bằng 0,4706. Toàn bộ 50 topic hội tụ. Summary dùng trung bình 99,82 từ, tương đương 99,82% ngân sách.")
    add_body(doc, "Recall thấp cần được diễn giải cùng ràng buộc độ dài. Reference DUC dài hơn prediction 100 từ, nên hệ thống không thể thu hồi toàn bộ câu tham chiếu. Notebook còn tính reference-sentence ceiling bằng cách chọn các câu reference ngắn nhất có thể nằm trong 100 từ; đây chỉ là upper bound về số câu, không phải upper bound chất lượng ngữ nghĩa.")
    add_heading(doc, "4.4. Phân phối similarity và lựa chọn threshold", level=2)
    add_body(doc, "P50 và P75 của similarity bằng 0; P90 là 0,0447 và P95 là 0,0762. Vì vậy threshold 0,10 loại nhiều cạnh có similarity cao hơn phần lớn phân phối. Kết quả sweep xác nhận threshold 0,02 là giá trị tốt nhất trong lưới đã thử, không phải một giá trị tối ưu phổ quát.")
    threshold_rows = [
        ["0,00", "0,1239", "0,0693", "1,0000", "0,0000", "292,68"],
        ["0,02", "0,1575", "0,0828", "0,1822", "0,0192", "51,32"],
        ["0,05", "0,1153", "0,0673", "0,0975", "0,0195", "27,26"],
        ["0,10", "0,0918", "0,0534", "0,0337", "0,0361", "9,42"],
        ["0,20", "0,0829", "0,0448", "0,0073", "0,3563", "2,04"],
        ["0,30", "0,0537", "0,0332", "0,0030", "0,6474", "0,85"],
    ]
    add_table(doc, ["τ", "Precision", "F1", "Density", "Isolated", "Avg. degree"], threshold_rows, [1.5, 2.8, 2.3, 2.6, 3.0, 3.5])
    add_caption(doc, "Bảng 4.2. Ảnh hưởng của ngưỡng similarity (các mốc tiêu biểu)")
    add_body(doc, "Khi τ = 0, đồ thị đầy đủ có density 1 và average degree 292,68; các cạnh bằng 0 vẫn được giữ theo điều kiện ≥ 0 nên cấu trúc trở nên quá dày. Khi τ tăng lên 0,30, isolated ratio đạt 0,6474 và average degree chỉ còn 0,85. Vùng 0,02 giữ mức liên kết đủ lớn nhưng không biến đồ thị thành đầy đủ, đồng thời đạt Precision và F1 cao nhất trong sweep.")
    add_heading(doc, "4.5. Thực nghiệm tinh chỉnh ngưỡng ở vùng thấp", level=2)
    add_body(doc, "Sau sweep thô, thí nghiệm bổ sung thu hẹp khoảng tìm kiếm về 0–0,03 với bước 0,0025. Trong từng nhánh, chỉ similarity threshold thay đổi; TF–IDF, damping 0,85, tolerance 10⁻⁸ và giới hạn hội tụ được giữ cố định. Mọi cấu hình được đo trên cùng 34 topic train có reference không rỗng.")
    add_body(doc, "Hai nhánh có hợp đồng đánh giá khác nhau và được báo cáo riêng. Package hiện tại lấy 15 câu PageRank cao nhất, khớp nội dung sau casefold và chuẩn hóa khoảng trắng. Pipeline notebook chọn tuần tự bằng MMR trong ngân sách 100 từ và khớp khóa (docid, num). Vì số câu dự đoán và cách nhận diện true positive khác nhau, trị số F1 giữa hai bảng không được dùng để kết luận pipeline nào tốt hơn.")
    top_k_rows = [
        ["0,0000", "0,1539", "0,1609", "0,1550", "4", "0,2286"],
        ["0,0100", "0,1518", "0,1571", "0,1521", "4", "0,2169"],
        ["0,0125", "0,1598", "0,1662", "0,1607", "4", "0,2072"],
        ["0,0175", "0,1591", "0,1646", "0,1597", "5", "0,1879"],
        ["0,0200", "0,1534", "0,1579", "0,1535", "5", "0,1790"],
        ["0,0300", "0,1448", "0,1472", "0,1440", "5", "0,1470"],
        ["0,1000", "0,1211", "0,1195", "0,1185", "5", "0,0328"],
    ]
    add_table(doc, ["τ", "Precision", "Recall", "F1", "F1=0", "Density"], top_k_rows, [1.8, 2.7, 2.5, 2.2, 2.0, 2.5])
    add_caption(doc, "Bảng 4.3. Sweep mịn PageRank + Top-K=15 trên train")
    add_body(doc, "Ở nhánh Top-K=15, τ=0,0125 đạt Macro F1 0,1607, cao hơn baseline τ=0,10 là 0,0422 điểm tuyệt đối, tương đương khoảng 35,6%. Precision và Recall cùng tăng, còn số topic F1 bằng 0 giảm từ 5 xuống 4. Density 0,2072, isolated ratio 0,0190, average degree 60,60 và trung bình 7,53 connected components cho thấy đồ thị vẫn đủ liên kết để truyền điểm.")
    mmr_fine_rows = [
        ["Train", "34", "0,1225", "0,0510", "0,0709", "13", "50/50"],
        ["Test", "9", "0,1611", "0,0652", "0,0889", "2", "9/9"],
    ]
    add_table(doc, ["Split", "Topic", "Precision", "Recall", "F1", "F1=0", "Hội tụ"], mmr_fine_rows, [2.2, 1.8, 2.5, 2.2, 2.0, 1.8, 2.2])
    add_caption(doc, "Bảng 4.4. Xác nhận τ=0,0125 trên pipeline MMR/100 từ")
    add_body(doc, "Trong sweep MMR/100 từ, F1 train cao nhất nằm tại τ=0,0275 với 0,0878; τ=0,0125 đạt 0,0709. Báo cáo vẫn khóa 0,0125 theo quyết định từ nhánh Top-K và đánh giá test đúng một lần, thay vì dùng kết quả test để đổi lại threshold. Do đó 0,0125 được gọi là cấu hình được chọn cho lần cập nhật, không phải ngưỡng tối ưu chung cho mọi chiến lược chọn câu.")
    fine_chart_dir = ROOT / "data" / "output" / "fine-threshold-sweep" / "charts"
    doc.add_picture(str(fine_chart_dir / "fine-threshold-quality.png"), width=Cm(14.2))
    add_caption(doc, "Hình 4.1. Chất lượng khi tinh chỉnh ngưỡng ở vùng thấp")
    doc.add_picture(str(fine_chart_dir / "fine-threshold-graph-health.png"), width=Cm(14.2))
    add_caption(doc, "Hình 4.2. Sức khỏe đồ thị khi tinh chỉnh ngưỡng")
    add_heading(doc, "4.6. Tham số PageRank và hội tụ", level=2)
    add_body(doc, "Notebook thử 45 tổ hợp gồm damping 0,70–0,95, tolerance 10⁻⁴–10⁻⁸ và max iterations 100–1000. Có 42/45 cấu hình hội tụ trên toàn bộ 50 topic. Nhóm tốt nhất dùng damping 0,85 và tolerance 10⁻⁸, đạt Precision 0,1575, F1 0,0828 và trung bình 39,64 vòng. Max iterations 100, 300 và 1000 tạo cùng metric trong nhóm này vì thuật toán thực tế dừng trước giới hạn.")
    add_body(doc, "Trong local grid cuối, 300 vòng được chọn thay cho 1000 khi các tiêu chí chất lượng và redundancy bằng nhau, đồng thời vẫn tạo khoảng an toàn lớn so với trung bình 39,64 vòng và hội tụ trên toàn train. Kết quả này minh họa rằng max_iterations không nên được hiểu là số vòng thực thi bắt buộc; nó chỉ là giới hạn trên.")
    add_heading(doc, "4.7. Ảnh hưởng của MMR", level=2)
    mmr_rows = [
        ["MMR off", "0,1350", "0,0733", "0,1390"],
        ["λ = 0,30", "0,1212", "0,0667", "0,0251"],
        ["λ = 0,50", "0,1529", "0,0800", "0,0509"],
        ["λ = 0,70", "0,1575", "0,0828", "0,0843"],
        ["λ = 0,90", "0,1360", "0,0738", "0,1251"],
        ["λ = 1,00", "0,1350", "0,0733", "0,1390"],
    ]
    add_table(doc, ["Thiết lập", "Precision", "F1", "Mean redundancy"], mmr_rows, [4.0, 3.5, 3.4, 4.8])
    add_caption(doc, "Bảng 4.5. So sánh các cấu hình MMR trên train")
    add_body(doc, "MMR off và λ = 1 cho cùng kết quả vì không có hình phạt redundancy. λ = 0,30 giảm mean redundancy xuống 0,0251 nhưng Precision cũng giảm còn 0,1212, cho thấy ưu tiên novelty quá mạnh đã loại các câu trung tâm. λ = 0,70 tạo điểm cân bằng tốt nhất trong lưới: Precision và F1 cao nhất, trong khi redundancy thấp hơn rõ rệt so với chọn thuần PageRank.")
    add_heading(doc, "4.8. Cấu hình cuối và so sánh train–test", level=2)
    add_body(doc, "Cấu hình cập nhật được khóa với threshold 0,0125; damping 0,85; tolerance 10⁻⁸; max iterations 300; budget 100 từ; MMR bật với λ=0,70. Cấu hình hội tụ trên toàn bộ 50 topic train và 9 topic test.")
    add_table(doc, ["Metric", "Train", "Test", "Diễn giải"], [
        ["Eligible topics", "34", "9", "Topic có reference hợp lệ"],
        ["Macro Precision", "0,1225", "0,1611", "Test cao hơn train"],
        ["Macro Recall", "0,0510", "0,0652", "Test cao hơn train"],
        ["Macro F1", "0,0709", "0,0889", "Test cao hơn train"],
        ["Hit@1", "0,2647", "0,1111", "Nhạy với câu đầu được chọn"],
        ["Hit@3", "0,3824", "0,6667", "Ít nhất một hit trong ba câu"],
        ["Hit@5", "0,6176", "0,7778", "Không tăng thêm trên test từ K=3"],
        ["Mean words", "99,88", "100,00", "Gần sử dụng hết budget"],
        ["Mean redundancy", "0,0825", "0,1050", "Test có độ trùng lặp cao hơn"],
        ["Mean PR iterations", "39,34", "39,33", "Hội tụ ổn định"],
    ], [4.2, 2.4, 2.4, 6.7])
    add_caption(doc, "Bảng 4.6. Kết quả cấu hình cuối trên train và test")
    add_body(doc, "Cả Precision, Recall và F1 đều cao hơn trên 9 topic test, nhưng tập test nhỏ nên chênh lệch này chưa đủ để kết luận cấu hình tổng quát tốt cho mọi dữ liệu. Bằng chứng chắc chắn hơn là toàn bộ topic hội tụ và cấu hình test không được dùng để điều chỉnh lại threshold.")
    add_heading(doc, "4.9. So sánh baseline và cấu hình được chọn", level=2)
    add_table(doc, ["Chỉ số train", "Baseline τ=0,10", "Cấu hình τ=0,0125", "Thay đổi"], [
        ["Macro Precision", "0,0918", "0,1225", "+0,0307"],
        ["Macro Recall", "0,0383", "0,0510", "+0,0127"],
        ["Macro F1", "0,0534", "0,0709", "+0,0175"],
        ["Hit@3", "0,3824", "0,3824", "0,0000"],
        ["Mean density", "0,0337", "0,2118", "+0,1781"],
        ["Mean isolated ratio", "0,0361", "0,0192", "−0,0169"],
    ], [4.1, 3.7, 4.1, 3.8])
    add_caption(doc, "Bảng 4.7. Kết quả baseline và cấu hình được chọn")
    add_body(doc, "Cải thiện chính đến từ việc hạ threshold để giữ thêm quan hệ giữa các câu. Đồ thị cuối dày hơn baseline nhưng vẫn xa đồ thị đầy đủ; tỷ lệ nút cô lập giảm. Điều này tạo đủ đường truyền điểm cho PageRank và giúp MMR lựa chọn trong một tập ứng viên có cấu trúc liên kết phong phú hơn.")
    add_heading(doc, "4.10. Error analysis và đe dọa đối với tính hợp lệ", level=2)
    add_bullets(doc, [
        "Exact (docid, num) match không ghi nhận câu diễn đạt tương đương hoặc cung cấp cùng thông tin bằng từ khác.",
        "Reference rỗng ở 16 topic train làm macro metric chỉ phản ánh 34 topic hợp lệ, dù convergence vẫn được kiểm tra trên đủ 50 topic.",
        "Precision ưu tiên câu trùng reference, nhưng không trực tiếp đo mạch lạc, ngữ pháp liên câu hoặc khả năng đọc độc lập.",
        "Recall bị giới hạn bởi budget 100 từ trong khi reference dài hơn; vì thế không nên so sánh Recall mà bỏ qua độ dài.",
        "TF–IDF dựa vào đồng hiện từ vựng, nên chưa xử lý đồng nghĩa, đa nghĩa, thực thể và đồng tham chiếu.",
        "MMR chỉ phạt similarity lớn nhất với câu đã chọn; nó chưa nhận biết hai câu khác từ nhưng cùng thông tin.",
        "Lưới tham số hữu hạn. Cấu hình được gọi là tốt nhất trong các giá trị đã thử, không phải tối ưu toàn cục.",
        "Thời gian chạy notebook là số đo trên môi trường hiện tại và không nên khái quát thành benchmark phần cứng khác.",
        "main.py đang return sau phép kiểm tra TF–IDF; phiên bản package chưa phải bằng chứng chạy end-to-end nếu chưa bỏ điểm dừng tạm thời.",
        "tests/test_main.py đang rỗng; phép đối chiếu sklearn hiện là kiểm tra thủ công, chưa phải regression test tự động.",
    ])
    add_heading(doc, "4.11. Đối chiếu với yêu cầu bài toán", level=2)
    add_table(doc, ["Yêu cầu", "Bằng chứng thực hiện"], [
        ["Đọc 50 tập tin", "Pipeline xử lý 50 topic train và tạo đúng một output cho mỗi topic"],
        ["Tách và biểu diễn câu", "preprocessing.py giữ metadata; tfidf.py viết tay và đối chiếu sklearn"],
        ["Biểu diễn thành đồ thị", "Đỉnh là câu, cạnh là cosine vượt threshold, trọng số bằng similarity"],
        ["Xếp hạng câu", "PageRank có damping, dangling redistribution và kiểm tra hội tụ"],
        ["Chọn thông tin quan trọng", "MMR kết hợp PageRank, không vượt ngân sách 100 từ"],
        ["Đánh giá và so sánh", "Sweep threshold/PageRank/MMR; macro metric; train/test tách biệt"],
        ["Phân tích ưu nhược điểm", "Mục 4.10 nêu giới hạn metric, dữ liệu, biểu diễn và trạng thái code"],
    ], [5.0, 10.7])
    add_caption(doc, "Bảng 4.8. Đối chiếu với yêu cầu bài toán")


def add_chapter_5_v2(doc: Document) -> None:
    add_heading(doc, "CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    add_heading(doc, "5.1. Kết quả đạt được", level=2)
    add_body(doc, "Đề tài đã có một pipeline notebook tóm tắt trích xuất có thể truy vết và đánh giá. Notebook đọc đủ 50 topic train, bảo toàn định danh câu, tự tính TF–IDF và cosine similarity, dựng đồ thị vô hướng có trọng số, cài đặt PageRank với xử lý dangling nodes, chọn câu bằng MMR trong ngân sách 100 từ và trả câu về thứ tự nguồn. Pipeline notebook tạo đúng một output không rỗng cho mỗi topic và mọi summary đều tuân thủ giới hạn từ.")
    add_body(doc, "Ở hướng phát triển mã nguồn, các bước đã được tách sang src/nlp_practice. TF–IDF viết tay được tinh chỉnh bằng chuẩn hóa L2 và được đối chiếu với TfidfVectorizer trên ba câu kiểm tra. Cả ba vector khớp từng trọng số, với sai khác lớn nhất bằng 0,0 trong phiên chạy Python 3.14.4 và scikit-learn 1.9.0. Kết quả này cung cấp bằng chứng trực tiếp rằng công thức TF dạng đếm, smooth IDF và L2 của code viết tay đồng nhất với cấu hình tham chiếu.")
    add_body(doc, "Bộ thực nghiệm biến threshold từ giá trị giả định thành lựa chọn dựa trên dữ liệu. Sweep mịn của PageRank + Top-K=15 chọn 0,0125 với Macro F1 0,1607. Khi giữ ngưỡng này cho pipeline MMR/100 từ, cấu hình đạt Macro Precision 0,1225 và F1 0,0709 trên train; trên test đạt Precision 0,1611 và F1 0,0889. Sweep MMR đạt F1 train cao hơn tại 0,0275, vì vậy 0,0125 chỉ được khẳng định là cấu hình đã khóa cho lần cập nhật, không phải tối ưu toàn cục.")
    add_heading(doc, "5.2. Những nội dung chưa giải quyết", level=2)
    add_body(doc, "Kết quả hiện tại chưa chứng minh chất lượng ngữ nghĩa hoặc tính mạch lạc của summary. Exact sentence match bỏ sót paraphrase và phụ thuộc vào cách chọn câu của reference. Hệ thống chưa có đánh giá con người, chưa so sánh với baseline vị trí như Lead, chưa dùng sentence embedding và chưa xử lý đồng tham chiếu. Mười sáu reference train rỗng làm giảm số topic có thể dùng cho metric chất lượng.")
    add_body(doc, "Về kỹ thuật, pairwise similarity vẫn có chi phí O(n²), không phù hợp trực tiếp với topic rất dài. Stopword list và tokenizer đang thiên về tiếng Anh và dựa trên luật. MMR dùng một λ cố định cho mọi topic. Cấu hình chỉ được chọn trong lưới hữu hạn; không nên diễn giải là tham số tối ưu cho dữ liệu ngoài DUC. Việc chuyển code sang package cũng chưa hoàn tất: selection.py chưa chứa toàn bộ MMR/budget, main.py dừng sau đối chiếu TF–IDF và tests/test_main.py chưa có regression test.")
    add_heading(doc, "5.3. Hướng cải tiến ưu tiên", level=2)
    add_bullets(doc, [
        "Bổ sung baseline Lead và random có kiểm soát để định lượng giá trị tăng thêm của TextRank.",
        "Thực hiện đánh giá con người theo mức đầy đủ, không dư thừa, mạch lạc và khả năng đọc độc lập.",
        "Bổ sung phép đối sánh nội dung hoặc metric ngữ nghĩa bên cạnh exact (docid, num), nhưng giữ nguyên phân tách train/test.",
        "Thử sentence embedding hoặc mô hình lai TF–IDF + embedding để nhận biết đồng nghĩa.",
        "Bổ sung đặc trưng vị trí, tiêu đề, thực thể và độ dài bằng personalized PageRank.",
        "Nhận biết câu phụ thuộc ngữ cảnh và có thể mở rộng lựa chọn bằng câu đứng trước.",
        "Dùng k-nearest-neighbor graph hoặc top-k cạnh cho mỗi câu để giảm bộ nhớ và thời gian.",
        "Phân tích metric theo độ dài topic, số nguồn và số thành phần liên thông để tìm nhóm thất bại.",
        "Khảo sát λ thích nghi theo redundancy của topic thay vì dùng một giá trị cố định.",
        "Lưu ranking đầy đủ trước MMR để có thể đánh giá thứ tự bằng MAP hoặc NDCG khi có ground truth phù hợp.",
        "Chuyển fixture TF–IDF thành unit test dùng math.isclose; bổ sung token một ký tự, Unicode, vector rỗng và sai khác tokenizer.",
        "Chuyển MMR/budget vào selection.py, bỏ return tạm trong main.py và chạy package end-to-end trên toàn bộ DUC.",
    ])
    add_heading(doc, "5.4. Kết luận chung", level=2)
    add_body(doc, "TextRank là baseline phù hợp cho 50 tập tin DUC vì không cần huấn luyện và có thể truy vết từ câu nguồn đến summary. TF–IDF viết tay khớp với scikit-learn, còn package src/nlp_practice tạo nền tảng cho kiểm thử và tái sử dụng. Bước tiếp theo là hoàn tất pipeline package end-to-end và tái tạo toàn bộ metric.")


def add_references(doc: Document) -> None:
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "[1] R. Mihalcea và P. Tarau, “TextRank: Bringing Order into Text,” Proceedings of EMNLP 2004, tr. 404–411. https://aclanthology.org/W04-3252/",
        "[2] L. Page, S. Brin, R. Motwani và T. Winograd, “The PageRank Citation Ranking: Bringing Order to the Web,” Stanford InfoLab Technical Report, 1999. https://ilpubs.stanford.edu/422/",
        "[3] C. D. Manning, P. Raghavan và H. Schütze, Introduction to Information Retrieval. Cambridge University Press, 2008. https://nlp.stanford.edu/IR-book/",
        "[4] NIST, “DUC 2004: Documents, Tasks, and Measures.” https://duc.nist.gov/duc2004/",
        "[5] J. Carbonell và J. Goldstein, “The Use of MMR, Diversity-Based Reranking for Reordering Documents and Producing Summaries,” SIGIR 1998.",
        "[6] Scikit-learn Developers, TfidfVectorizer, scikit-learn 1.9 documentation.",
        "[7] Mã nguồn dự án, src/nlp_practice/tfidf.py và src/nlp_practice/main.py, phiên bản kiểm tra ngày 28/08/2026.",
        "[8] Notebook dự án, duc-textrank-pipeline.ipynb, kết quả thực thi ngày 20/08/2026.",
        "[9] Notebook dự án, duc-textrank-evaluation.ipynb, kết quả thực thi ngày 20/08/2026.",
        "[10] Bộ notebook dự án, textrank-parameter-analysis/01–05, kết quả thực thi ngày 20/08/2026.",
        "[11] Script và artifact dự án, notebooks/textrank-parameter-analysis/fine_threshold_sweep.py và data/output/fine-threshold-sweep/, kết quả ngày 28/08/2026.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style="Normal")
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.25
        for run in p.runs:
            run.font.size = Pt(11.5)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    actual_sha = hashlib.sha256(REFERENCE.read_bytes()).hexdigest()
    if actual_sha != EXPECTED_SHA:
        raise SystemExit(f"Template SHA mismatch: {actual_sha}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    configure_styles(doc)
    replace_cover_topic(doc)
    rebuild_front_lists(doc)

    # The retained template contains a sample acronym table for another topic.
    # Remove it before adding this report's own data tables.
    for table in list(doc.tables):
        table._element.getparent().remove(table._element)

    # Remove the sample academic outline while preserving all front matter.
    started = False
    for p in list(doc.paragraphs):
        if p.text.strip() == "CHƯƠNG 1. TỔNG QUAN":
            started = True
        if started:
            delete_paragraph(p)

    # Make content start on a new page but keep the template's section/footer system.
    doc.add_page_break()
    figs = make_figures()
    add_chapter_1_v2(doc)
    add_chapter_2_v2(doc)
    add_chapter_3_v2(doc, figs)
    add_chapter_4_v2(doc)
    add_chapter_5_v2(doc)
    add_references(doc)
    set_update_fields(doc)
    doc.core_properties.title = "Tóm tắt trích xuất DUC bằng TextRank/PageRank và MMR"
    doc.core_properties.subject = "Báo cáo môn Khai thác thông tin - HUTECH"
    doc.core_properties.author = "Học viên thực hiện"
    doc.core_properties.keywords = "TextRank, PageRank, MMR, extractive summarization, TF-IDF, exact sentence match, DUC"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
